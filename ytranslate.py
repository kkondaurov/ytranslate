#!/usr/bin/env python3
import argparse
import hashlib
import json
import math
import os
import random
import re
import shutil
import sys
import subprocess
import time
import unicodedata
import inspect
from concurrent.futures import ThreadPoolExecutor, as_completed
from typing import Any, Callable, Dict, List, Optional, Tuple
from urllib.parse import parse_qs, urlparse

import requests
from openai import OpenAI
import openai
from docx import Document
from docx.shared import Pt
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import NoTranscriptFound, TranscriptsDisabled


YOUTUBE_API_URL = "https://www.googleapis.com/youtube/v3/videos"
PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
ENV_PATH = os.path.join(PROJECT_ROOT, ".env")
DEFAULT_MODEL = "gpt-5.4-mini"
OPENAI_ASR_MODEL = "gpt-4o-transcribe-diarize"
OPENAI_TRANSCRIBE_URL = "https://api.openai.com/v1/audio/transcriptions"
DEFAULT_TARGET_LANGUAGE = "Russian"
OPENAI_TIMEOUT_SECONDS = 1800
OPENAI_ASR_TIMEOUT_SECONDS = 600
OPENAI_TEMPERATURE = 0.2
OPENAI_CLEANUP_TEMPERATURE = 0.0
OPENAI_ANNOTATION_TEMPERATURE = 0.0
ASR_MAX_CHUNK_SECONDS = 1400
ASR_CHUNK_SECONDS = 600
ASR_AUDIO_BITRATE = "64k"
ASR_SAMPLE_RATE = "16000"
ASR_MAX_UPLOAD_BYTES = 24 * 1024 * 1024
ASR_JOBS = 1
ASR_MAX_RETRIES = 3
ASR_MAX_PASSES = 3
ASR_RETRY_PASS_DELAY_SECONDS = 30
OPENAI_ASR_TRANSPORT = "curl"
DOCX_FONT_NAME = "Arial"
DOCX_FONT_SIZE = Pt(13)
DOCX_HEADING_FONT_SIZE = Pt(16)
OUTPUT_DIR = os.path.expanduser(os.getenv("YTRANSLATE_OUTPUT_DIR", "~/Downloads"))
CACHE_DIR = os.path.expanduser("~/Library/Caches/ytranslate")
SPEAKER_OVERRIDES_FILENAME = "speaker-overrides.json"
TURN_TEXT_PASS_MAX_CHARS = 6_000
VOICE_RECONCILIATION_MIN_SEGMENT_SECONDS = 2.4
VOICE_RECONCILIATION_MAX_SEGMENT_SECONDS = 16.0
VOICE_RECONCILIATION_MIN_SIMILARITY = 0.86
VOICE_RECONCILIATION_MIN_MARGIN = 0.035
VOICE_RECONCILIATION_NEIGHBOR_GAP_SECONDS = 1.0
VOICE_RECONCILIATION_MIN_ANCHOR_SEGMENTS = 2
VOICE_RECONCILIATION_MIN_ANCHOR_SECONDS = 5.0
VOICE_RECONCILIATION_MAX_ANCHOR_SEGMENTS = 28
VOICE_RECONCILIATION_LOCAL_MAJORITY_SHARE = 0.78
VOICE_RECONCILIATION_LOCAL_MAJORITY_MIN_SEGMENTS = 2
SPEAKER_IDENTITY_LINKER_MODEL = os.getenv(
    "YTRANSLATE_SPEAKER_IDENTITY_LINKER_MODEL", "gpt-5.6-luna"
)
SPEAKER_IDENTITY_LINKER_REASONING_EFFORT = "low"
SPEAKER_IDENTITY_LINKER_BATCH_SEGMENTS = 300
SPEAKER_IDENTITY_LINKER_CONTEXT_SEGMENTS = 12
SPEAKER_IDENTITY_LINKER_CACHE_SCHEMA_VERSION = 1
ROLE_SPEAKER_LABEL_MARKERS = (
    "ad read",
    "advertisement",
    "sponsor read",
    "sponsor",
    "commercial",
)
VOICE_DERIVED_SPEAKER_SOURCES = {"voice", "voice_neighbor", "voice_local_majority"}
ALL_IN_CHANNEL_TITLES = {"all-in podcast", "all in podcast"}
ALL_IN_KNOWN_SPEAKERS = [
    {
        "id": "speaker_jason_calacanis",
        "label_short": "Jason",
        "label_full": "Jason Calacanis",
        "aliases": ["jason", "j-cal", "jcal", "j cal"],
    },
    {
        "id": "speaker_chamath_palihapitiya",
        "label_short": "Chamath",
        "label_full": "Chamath Palihapitiya",
        "aliases": ["chamath", "chumath", "jamath"],
    },
    {
        "id": "speaker_david_sacks",
        "label_short": "Sacks",
        "label_full": "David Sacks",
        "aliases": ["sacks", "zach", "sachs", "david sacks"],
    },
    {
        "id": "speaker_david_friedberg",
        "label_short": "Friedberg",
        "label_full": "David Friedberg",
        "aliases": ["friedberg", "freeberg", "freiberg", "david friedberg"],
    },
]


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description=(
            "Translate a YouTube video's transcript into a target language "
            "and structure it as a conversation."
        )
    )
    parser.add_argument("url", help="YouTube video URL")
    parser.add_argument(
        "target_language",
        nargs="?",
        help="Target language (defaults to DEFAULT_TARGET_LANGUAGE or Russian)",
    )
    parser.add_argument(
        "--docx-test",
        action="store_true",
        help="Generate a sample DOCX without calling external APIs",
    )
    parser.add_argument(
        "--debug",
        action="store_true",
        help="Write per-stage debug artifacts as Markdown/JSON instead of DOCX/PDF",
    )
    return parser.parse_args()

def load_dotenv(path: str) -> None:
    if not os.path.exists(path):
        return
    try:
        with open(path, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                key, value = line.split("=", 1)
                key = key.strip()
                value = value.strip().strip("\"").strip("'")
                if key and key not in os.environ:
                    os.environ[key] = value
    except OSError:
        return


def load_project_env() -> None:
    load_dotenv(ENV_PATH)


def resolve_target_language(target_language: Optional[str]) -> str:
    value = (target_language or os.getenv("DEFAULT_TARGET_LANGUAGE") or DEFAULT_TARGET_LANGUAGE).strip()
    return value or DEFAULT_TARGET_LANGUAGE


def extract_video_id(url: str) -> Optional[str]:
    if not url:
        return None
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    path = parsed.path
    query = parse_qs(parsed.query)

    vid = query.get("v", [None])[0]
    if vid:
        return vid

    if host.endswith("youtu.be"):
        vid = path.strip("/").split("/")[0]
        return vid or None

    if "youtube.com" in host:
        if path.startswith("/watch"):
            return query.get("v", [None])[0]
        if path.startswith("/shorts/"):
            parts = path.split("/")
            return parts[2] if len(parts) > 2 else None
        if path.startswith("/embed/"):
            parts = path.split("/")
            return parts[2] if len(parts) > 2 else None
        if path.startswith("/live/"):
            parts = path.split("/")
            return parts[2] if len(parts) > 2 else None

    return None


def canonicalize_youtube_url(url: str) -> Optional[str]:
    video_id = extract_video_id(url)
    if not video_id:
        return None
    return f"https://youtu.be/{video_id}"


def metadata_from_youtube_item(item: Dict[str, Any]) -> Dict[str, Any]:
    snippet = item.get("snippet", {})
    tags = snippet.get("tags") or []
    if not isinstance(tags, list):
        tags = []
    return {
        "title": snippet.get("title", ""),
        "description": snippet.get("description", ""),
        "channelTitle": snippet.get("channelTitle"),
        "channelId": snippet.get("channelId"),
        "tags": [str(tag) for tag in tags],
        "defaultLanguage": snippet.get("defaultLanguage"),
        "defaultAudioLanguage": snippet.get("defaultAudioLanguage"),
    }


def fetch_video_metadata(video_id: str, api_key: str) -> Dict[str, Any]:
    params = {
        "part": "snippet",
        "id": video_id,
        "key": api_key,
    }
    resp = requests.get(YOUTUBE_API_URL, params=params, timeout=30)
    if resp.status_code != 200:
        msg = ""
        try:
            data = resp.json()
            msg = data.get("error", {}).get("message", "")
        except Exception:
            msg = resp.text
        raise RuntimeError(
            f"YouTube API error ({resp.status_code}): {msg or 'Unknown error'}"
        )

    data = resp.json()
    items = data.get("items", [])
    if not items:
        raise RuntimeError("No video metadata found (invalid video ID?)")

    return metadata_from_youtube_item(items[0])


def normalize_identity_text(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value or "")
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii").lower()
    return re.sub(r"[^a-z0-9]+", " ", ascii_text).strip()


def speaker_public_fields(speaker: Dict[str, Any]) -> Dict[str, str]:
    speaker_id = str(speaker.get("id") or speaker_id_from_label(str(speaker.get("label_short") or "speaker")))
    label_short = str(speaker.get("label_short") or speaker_id)
    label_full = str(speaker.get("label_full") or label_short)
    return {
        "id": speaker_id,
        "label_short": label_short,
        "label_full": label_full,
    }


def infer_known_speaker_roster(metadata: Dict[str, Any]) -> List[Dict[str, Any]]:
    channel_title = normalize_identity_text(str(metadata.get("channelTitle") or ""))
    tags = [normalize_identity_text(str(tag)) for tag in metadata.get("tags", [])]
    tag_text = " ".join(tags)
    is_all_in = (
        channel_title in ALL_IN_CHANNEL_TITLES
        or "all in podcast" in tag_text
        or "jason calacanis" in tag_text
        and "chamath" in tag_text
        and "david sacks" in tag_text
        and "david friedberg" in tag_text
    )
    if not is_all_in:
        return []
    return [dict(speaker) for speaker in ALL_IN_KNOWN_SPEAKERS]


def pick_transcript(transcripts: List[Any], preferred_langs: List[str]) -> Any:
    def find_match(predicate):
        for t in transcripts:
            if predicate(t):
                return t
        return None

    for lang in preferred_langs:
        if not lang:
            continue
        match = find_match(
            lambda t, lang=lang: (not t.is_generated) and t.language_code == lang
        )
        if match:
            return match

    manual = find_match(lambda t: not t.is_generated)
    if manual:
        return manual

    for lang in preferred_langs:
        if not lang:
            continue
        match = find_match(lambda t, lang=lang: t.language_code == lang)
        if match:
            return match

    return transcripts[0] if transcripts else None


def list_transcripts(video_id: str):
    if hasattr(YouTubeTranscriptApi, "list_transcripts"):
        return YouTubeTranscriptApi.list_transcripts(video_id)

    api = YouTubeTranscriptApi()
    if hasattr(api, "list"):
        return api.list(video_id)
    if hasattr(api, "list_transcripts"):
        return api.list_transcripts(video_id)

    raise RuntimeError("Unsupported youtube-transcript-api version")


def fetch_transcript(video_id: str, preferred_langs: List[str]) -> Dict[str, Any]:
    transcript_list = list_transcripts(video_id)
    transcripts = list(transcript_list)
    if not transcripts:
        raise NoTranscriptFound(video_id)

    transcript = pick_transcript(transcripts, preferred_langs)
    if not transcript:
        raise NoTranscriptFound(video_id)

    fetched = transcript.fetch()
    if hasattr(fetched, "to_raw_data"):
        segments = fetched.to_raw_data()
    else:
        segments = fetched

    return {
        "language_code": transcript.language_code,
        "language": transcript.language,
        "is_generated": transcript.is_generated,
        "segments": segments,
    }


def clean_segment_text(text: str) -> str:
    text = (text or "").replace("\n", " ").strip()
    text = re.sub(r"\s+", " ", text)
    return text


def normalize_segments(segments: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    normalized: List[Dict[str, Any]] = []
    last_text = None
    for seg in segments:
        if isinstance(seg, dict):
            text = clean_segment_text(seg.get("text", ""))
            start = seg.get("start")
            duration = seg.get("duration")
        else:
            text = clean_segment_text(getattr(seg, "text", ""))
            start = getattr(seg, "start", None)
            duration = getattr(seg, "duration", None)
        if not text:
            continue
        if text == last_text:
            continue
        normalized.append({
            "start": start,
            "duration": duration,
            "text": text,
        })
        last_text = text
    return normalized


SPEAKER_LABEL_RE = re.compile(
    r"^\s*(?:\[(?P<bracket>[^\]]{1,40})\]|(?P<plain>[A-Z][A-Za-z0-9 ._'’-]{0,40})):\s+(?P<body>\S.*)$"
)


def split_speaker_label(text: str) -> Optional[Dict[str, str]]:
    match = SPEAKER_LABEL_RE.match(text or "")
    if not match:
        return None
    label = (match.group("bracket") or match.group("plain") or "").strip()
    body = (match.group("body") or "").strip()
    if not label or not body:
        return None
    return {"label": label, "text": body}


def speaker_id_from_label(label: str) -> str:
    normalized = unicodedata.normalize("NFKD", label or "")
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii").lower()
    ascii_text = re.sub(r"[^a-z0-9]+", "_", ascii_text).strip("_")
    return f"speaker_{ascii_text or 'unknown'}"


def is_high_quality_youtube_transcript(transcript_info: Dict[str, Any]) -> bool:
    if transcript_info.get("is_generated"):
        return False
    segments = normalize_segments(transcript_info.get("segments", []))
    if not segments:
        return False
    labels = []
    for segment in segments:
        parsed = split_speaker_label(segment.get("text", ""))
        if parsed:
            labels.append(parsed["label"].lower())
    min_labeled = min(8, max(4, len(segments) // 10))
    return (
        len(set(labels)) >= 2
        and len(labels) >= min_labeled
        and (len(labels) / len(segments)) >= 0.3
    )


def make_speaker(label: str) -> Dict[str, str]:
    return {
        "id": speaker_id_from_label(label),
        "label_short": label,
        "label_full": label,
    }


def attributed_turns_from_labeled_segments(segments: List[Dict[str, Any]]) -> Dict[str, Any]:
    speakers_by_id: Dict[str, Dict[str, str]] = {}
    turns: List[Dict[str, str]] = []
    current_speaker_id: Optional[str] = None

    for segment in segments:
        text = clean_segment_text(segment.get("text", ""))
        if not text:
            continue
        parsed = split_speaker_label(text)
        if parsed:
            speaker = make_speaker(parsed["label"])
            speakers_by_id.setdefault(speaker["id"], speaker)
            current_speaker_id = speaker["id"]
            text = parsed["text"]
        if not current_speaker_id:
            current_speaker_id = "speaker_unknown"
            speakers_by_id.setdefault(
                current_speaker_id,
                {
                    "id": current_speaker_id,
                    "label_short": "Speaker",
                    "label_full": "Speaker",
                },
            )
        if turns and turns[-1].get("speaker_id") == current_speaker_id:
            turns[-1]["text_source"] = (turns[-1].get("text_source", "") + " " + text).strip()
        else:
            turns.append({"speaker_id": current_speaker_id, "text_source": text})

    return {
        "speakers": list(speakers_by_id.values()),
        "turns": turns,
    }


def get_ffmpeg_executable() -> str:
    system_ffmpeg = shutil.which("ffmpeg")
    if system_ffmpeg:
        return system_ffmpeg
    try:
        import imageio_ffmpeg
    except ImportError as exc:
        raise RuntimeError(
            "ffmpeg is required for OpenAI ASR. Install ffmpeg or install the "
            "imageio-ffmpeg Python package from requirements.txt."
        ) from exc
    return imageio_ffmpeg.get_ffmpeg_exe()


def get_video_cache_dir(video_id: str) -> str:
    cache_dir = os.path.join(CACHE_DIR, sanitize_filename(video_id))
    os.makedirs(cache_dir, exist_ok=True)
    return cache_dir


def get_speaker_mapping_override_path(video_id: str) -> str:
    return os.path.join(get_video_cache_dir(video_id), SPEAKER_OVERRIDES_FILENAME)


def download_youtube_audio(url: str, video_id: str, log: Callable[[str], None]) -> str:
    audio_dir = os.path.join(get_video_cache_dir(video_id), "audio")
    os.makedirs(audio_dir, exist_ok=True)
    for filename in sorted(os.listdir(audio_dir)):
        if filename.startswith("source.") and not filename.endswith(".part"):
            return os.path.join(audio_dir, filename)

    output_template = os.path.join(audio_dir, "source.%(ext)s")
    command = [
        sys.executable,
        "-m",
        "yt_dlp",
        "-f",
        "ba[ext=m4a]/ba[ext=webm]/ba/bestaudio",
        "-o",
        output_template,
        "--no-playlist",
        url,
    ]
    log("Downloading audio for OpenAI ASR...")
    try:
        subprocess.run(command, check=True)
    except subprocess.CalledProcessError as exc:
        raise RuntimeError("Failed to download YouTube audio with yt-dlp") from exc

    for filename in sorted(os.listdir(audio_dir)):
        if filename.startswith("source.") and not filename.endswith(".part"):
            return os.path.join(audio_dir, filename)
    raise RuntimeError("yt-dlp did not produce an audio file")


def transcode_and_chunk_audio(
    source_audio_path: str,
    video_id: str,
    chunk_seconds: int,
    log: Callable[[str], None],
) -> List[str]:
    chunk_dir = os.path.join(get_video_cache_dir(video_id), f"chunks-{chunk_seconds}s")
    os.makedirs(chunk_dir, exist_ok=True)
    existing = sorted(
        os.path.join(chunk_dir, name)
        for name in os.listdir(chunk_dir)
        if re.match(r"chunk-\d{3}\.mp3$", name)
    )
    if existing:
        return existing

    ffmpeg = get_ffmpeg_executable()
    output_pattern = os.path.join(chunk_dir, "chunk-%03d.mp3")
    command = [
        ffmpeg,
        "-hide_banner",
        "-y",
        "-i",
        source_audio_path,
        "-vn",
        "-ac",
        "1",
        "-ar",
        ASR_SAMPLE_RATE,
        "-b:a",
        ASR_AUDIO_BITRATE,
        "-f",
        "segment",
        "-segment_time",
        str(chunk_seconds),
        "-reset_timestamps",
        "1",
        output_pattern,
    ]
    log("Compressing and chunking audio for OpenAI ASR...")
    subprocess.run(command, check=True)
    chunks = sorted(
        os.path.join(chunk_dir, name)
        for name in os.listdir(chunk_dir)
        if re.match(r"chunk-\d{3}\.mp3$", name)
    )
    oversized = [path for path in chunks if os.path.getsize(path) > ASR_MAX_UPLOAD_BYTES]
    if oversized:
        raise RuntimeError(
            "Audio chunk exceeds OpenAI upload limit: "
            + ", ".join(os.path.basename(path) for path in oversized)
        )
    if not chunks:
        raise RuntimeError("ffmpeg did not produce any audio chunks")
    return chunks


def probe_audio_duration_seconds(audio_path: str) -> float:
    ffmpeg = get_ffmpeg_executable()
    completed = subprocess.run(
        [ffmpeg, "-hide_banner", "-i", audio_path, "-f", "null", "-"],
        check=False,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
    )
    match = re.search(r"Duration:\s*(\d+):(\d+):(\d+(?:\.\d+)?)", completed.stderr)
    if not match:
        raise RuntimeError(f"Unable to parse audio duration for {audio_path}")
    return int(match.group(1)) * 3600 + int(match.group(2)) * 60 + float(match.group(3))


def build_chunk_offsets(chunks: List[str]) -> List[float]:
    offsets: List[float] = []
    current = 0.0
    for chunk in chunks:
        offsets.append(round(current, 3))
        current += probe_audio_duration_seconds(chunk)
    return offsets


def get_asr_chunk_seconds() -> int:
    raw = os.getenv("OPENAI_ASR_CHUNK_SECONDS", str(ASR_CHUNK_SECONDS))
    try:
        chunk_seconds = int(raw)
    except ValueError as exc:
        raise RuntimeError(f"OPENAI_ASR_CHUNK_SECONDS must be an integer, got {raw!r}") from exc
    if chunk_seconds <= 0:
        raise RuntimeError("OPENAI_ASR_CHUNK_SECONDS must be positive")
    if chunk_seconds > ASR_MAX_CHUNK_SECONDS:
        raise RuntimeError(
            "OPENAI_ASR_CHUNK_SECONDS exceeds the OpenAI ASR model limit: "
            f"{chunk_seconds}s > {ASR_MAX_CHUNK_SECONDS}s. Use {ASR_CHUNK_SECONDS}s or lower."
        )
    return chunk_seconds


def get_asr_jobs() -> int:
    raw = os.getenv("OPENAI_ASR_JOBS", str(ASR_JOBS))
    try:
        jobs = int(raw)
    except ValueError as exc:
        raise RuntimeError(f"OPENAI_ASR_JOBS must be an integer, got {raw!r}") from exc
    if jobs <= 0:
        raise RuntimeError("OPENAI_ASR_JOBS must be positive")
    return jobs


def get_asr_timeout_seconds() -> int:
    raw = os.getenv("OPENAI_ASR_TIMEOUT_SECONDS", str(OPENAI_ASR_TIMEOUT_SECONDS))
    try:
        timeout_seconds = int(raw)
    except ValueError as exc:
        raise RuntimeError(f"OPENAI_ASR_TIMEOUT_SECONDS must be an integer, got {raw!r}") from exc
    if timeout_seconds <= 0:
        raise RuntimeError("OPENAI_ASR_TIMEOUT_SECONDS must be positive")
    return timeout_seconds


def get_asr_max_retries() -> int:
    raw = os.getenv("OPENAI_ASR_MAX_RETRIES", str(ASR_MAX_RETRIES))
    try:
        retries = int(raw)
    except ValueError as exc:
        raise RuntimeError(f"OPENAI_ASR_MAX_RETRIES must be an integer, got {raw!r}") from exc
    if retries <= 0:
        raise RuntimeError("OPENAI_ASR_MAX_RETRIES must be positive")
    return retries


def get_asr_max_passes() -> int:
    raw = os.getenv("OPENAI_ASR_MAX_PASSES", str(ASR_MAX_PASSES))
    try:
        passes = int(raw)
    except ValueError as exc:
        raise RuntimeError(f"OPENAI_ASR_MAX_PASSES must be an integer, got {raw!r}") from exc
    if passes <= 0:
        raise RuntimeError("OPENAI_ASR_MAX_PASSES must be positive")
    return passes


def get_asr_retry_pass_delay_seconds() -> int:
    raw = os.getenv("OPENAI_ASR_RETRY_PASS_DELAY_SECONDS", str(ASR_RETRY_PASS_DELAY_SECONDS))
    try:
        delay_seconds = int(raw)
    except ValueError as exc:
        raise RuntimeError(
            f"OPENAI_ASR_RETRY_PASS_DELAY_SECONDS must be an integer, got {raw!r}"
        ) from exc
    if delay_seconds < 0:
        raise RuntimeError("OPENAI_ASR_RETRY_PASS_DELAY_SECONDS must not be negative")
    return delay_seconds


def get_asr_transport() -> str:
    transport = os.getenv("OPENAI_ASR_TRANSPORT", OPENAI_ASR_TRANSPORT).strip().lower()
    if transport not in {"curl", "requests"}:
        raise RuntimeError("OPENAI_ASR_TRANSPORT must be either 'curl' or 'requests'")
    if transport == "curl" and not shutil.which("curl"):
        return "requests"
    return transport


def curl_config_quote(value: str) -> str:
    escaped = value.replace("\\", "\\\\").replace('"', '\\"').replace("\n", "")
    return f'"{escaped}"'


def transcribe_audio_chunk_with_curl(
    chunk_path: str,
    openai_key: str,
    asr_model: str,
    timeout_seconds: int,
) -> Dict[str, Any]:
    curl = shutil.which("curl")
    if not curl:
        raise RuntimeError("curl executable not found for OpenAI ASR upload")
    chunk_name = os.path.basename(chunk_path)
    form_file = f"file=@{chunk_path};type=audio/mpeg;filename={chunk_name}"
    config = "\n".join(
        [
            f"url = {curl_config_quote(OPENAI_TRANSCRIBE_URL)}",
            "request = POST",
            f"header = {curl_config_quote(f'Authorization: Bearer {openai_key}')}",
            f"form = {curl_config_quote(f'model={asr_model}')}",
            'form = "response_format=diarized_json"',
            'form = "chunking_strategy=auto"',
            f"form = {curl_config_quote(form_file)}",
            f"max-time = {timeout_seconds}",
            "connect-timeout = 30",
            "retry = 2",
            "retry-delay = 2",
            "retry-all-errors",
            "http1.1",
            "silent",
            "show-error",
            "fail-with-body",
        ]
    )
    completed = subprocess.run(
        [curl, "--config", "-"],
        input=config,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        text=True,
        timeout=timeout_seconds + 45,
        check=False,
    )
    if completed.returncode != 0:
        detail = (completed.stderr or completed.stdout or "").strip()
        raise RuntimeError(f"curl ASR upload failed for {chunk_name}: {detail[:1000]}")
    try:
        return json.loads(completed.stdout)
    except json.JSONDecodeError as exc:
        raise RuntimeError(
            f"curl ASR upload returned invalid JSON for {chunk_name}: {completed.stdout[:1000]}"
        ) from exc


def transcribe_audio_chunk(
    chunk_path: str,
    openai_key: str,
    asr_model: str,
    timeout_seconds: int,
    log: Optional[Callable[[str], None]] = None,
) -> Dict[str, Any]:
    headers = {"Authorization": f"Bearer {openai_key}"}
    data = {
        "model": asr_model,
        "response_format": "diarized_json",
        "chunking_strategy": "auto",
    }
    chunk_name = os.path.basename(chunk_path)
    max_retries = get_asr_max_retries()
    transport = get_asr_transport()
    for attempt in range(1, max_retries + 1):
        if log:
            log(
                f"OpenAI ASR request attempt {attempt}/{max_retries} "
                f"for {chunk_name} via {transport} (timeout={timeout_seconds}s)"
            )
        try:
            if transport == "curl":
                return transcribe_audio_chunk_with_curl(
                    chunk_path,
                    openai_key,
                    asr_model,
                    timeout_seconds,
                )
            with open(chunk_path, "rb") as audio_file:
                files = {"file": (chunk_name, audio_file, "audio/mpeg")}
                response = requests.post(
                    OPENAI_TRANSCRIBE_URL,
                    headers=headers,
                    data=data,
                    files=files,
                    timeout=timeout_seconds,
                )
        except (requests.RequestException, RuntimeError, subprocess.SubprocessError) as exc:
            if attempt == max_retries:
                raise RuntimeError(
                    f"OpenAI ASR upload failed after retries for {chunk_name}: {exc}"
                ) from exc
            if log:
                log(
                    f"OpenAI ASR retryable error on attempt {attempt}/{max_retries} "
                    f"for {chunk_name}: {exc}"
                )
            time.sleep(min(60, 5 * attempt))
            continue

        if response.status_code < 500 and response.status_code != 429:
            if response.status_code >= 400:
                raise RuntimeError(
                    f"OpenAI ASR failed for {chunk_name}: "
                    f"{response.status_code} {response.text[:1000]}"
                )
            return response.json()

        if attempt == max_retries:
            raise RuntimeError(
                f"OpenAI ASR failed after retries for {chunk_name}: "
                f"{response.status_code} {response.text[:1000]}"
            )
        if log:
            log(
                f"OpenAI ASR retryable response on attempt {attempt}/{max_retries} "
                f"for {chunk_name}: {response.status_code}"
            )
        time.sleep(min(60, 5 * attempt))

    raise RuntimeError("Unreachable OpenAI ASR retry state")


def extract_diarized_segments(response: Dict[str, Any]) -> List[Dict[str, Any]]:
    segments = response.get("segments") or []
    normalized: List[Dict[str, Any]] = []
    for segment in segments:
        text = clean_segment_text(str(segment.get("text") or ""))
        if not text:
            continue
        start = float(segment.get("start") or 0)
        end = float(segment.get("end") or start)
        normalized.append(
            {
                "speaker": str(segment.get("speaker") or "speaker"),
                "start": start,
                "end": end,
                "text": text,
            }
        )
    if normalized:
        return normalized
    text = clean_segment_text(str(response.get("text") or ""))
    return [{"speaker": "speaker", "start": 0.0, "end": 0.0, "text": text}] if text else []


def merge_diarized_chunks(chunks: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    merged: List[Dict[str, Any]] = []
    for chunk in chunks:
        offset = float(chunk.get("offset_seconds") or 0)
        chunk_index = int(chunk.get("chunk_index") or 0)
        chunk_name = str(chunk.get("chunk") or "")
        for segment in chunk.get("segments", []):
            text = clean_segment_text(segment.get("text", ""))
            if not text:
                continue
            start = float(segment.get("start") or 0) + offset
            end = float(segment.get("end") or segment.get("start") or 0) + offset
            local_speaker = str(segment.get("speaker") or "speaker")
            merged.append(
                {
                    "speaker": local_speaker,
                    "local_speaker": local_speaker,
                    "chunk_index": chunk_index,
                    "chunk": chunk_name,
                    "start": round(start, 3),
                    "end": round(end, 3),
                    "text": text,
                }
            )
    return merged


def build_local_speaker_profiles(
    segments: List[Dict[str, Any]],
    max_chars_per_profile: int = 1800,
) -> List[Dict[str, Any]]:
    profiles_by_key: Dict[Any, Dict[str, Any]] = {}
    order: List[Any] = []
    for segment in segments:
        chunk_index = int(segment.get("chunk_index") or 0)
        local_speaker = str(segment.get("local_speaker") or segment.get("speaker") or "speaker")
        key = (chunk_index, local_speaker)
        if key not in profiles_by_key:
            profiles_by_key[key] = {
                "chunk_index": chunk_index,
                "local_speaker": local_speaker,
                "segment_count": 0,
                "start": float(segment.get("start") or 0),
                "end": float(segment.get("end") or segment.get("start") or 0),
                "samples": [],
                "_sample_chars": 0,
            }
            order.append(key)
        profile = profiles_by_key[key]
        profile["segment_count"] += 1
        profile["end"] = float(segment.get("end") or segment.get("start") or profile["end"])
        text = clean_segment_text(segment.get("text", ""))
        if text and profile["_sample_chars"] < max_chars_per_profile:
            remaining = max_chars_per_profile - profile["_sample_chars"]
            sample = text[:remaining]
            profile["samples"].append(sample)
            profile["_sample_chars"] += len(sample)

    profiles: List[Dict[str, Any]] = []
    for key in order:
        profile = dict(profiles_by_key[key])
        profile.pop("_sample_chars", None)
        profile["start"] = round(float(profile["start"]), 3)
        profile["end"] = round(float(profile["end"]), 3)
        profiles.append(profile)
    return profiles


def roster_speaker_by_alias(roster: List[Dict[str, Any]]) -> Dict[str, Dict[str, Any]]:
    by_alias: Dict[str, Dict[str, Any]] = {}
    for speaker in roster:
        aliases = [
            speaker.get("label_short", ""),
            speaker.get("label_full", ""),
            *speaker.get("aliases", []),
        ]
        for alias in aliases:
            normalized = normalize_identity_text(str(alias))
            if normalized:
                by_alias[normalized] = speaker
    return by_alias


def find_roster_mentions(text: str, roster: List[Dict[str, Any]]) -> List[Tuple[str, Dict[str, Any]]]:
    normalized_text = f" {normalize_identity_text(text)} "
    matches: List[Tuple[str, Dict[str, Any]]] = []
    for alias, speaker in sorted(
        roster_speaker_by_alias(roster).items(),
        key=lambda item: len(item[0]),
        reverse=True,
    ):
        if f" {alias} " in normalized_text:
            matches.append((alias, speaker))
    return matches


def is_likely_handoff_phrase(text: str, alias: str) -> bool:
    normalized = normalize_identity_text(text)
    if alias not in normalized:
        return False
    handoff_markers = [
        "what s your take",
        "whats your take",
        "what is your take",
        "what do you think",
        "your thoughts",
        "do you worry",
        "running your",
        "curious to get your take",
        "you want to add",
        "as we wrap",
        "how s this",
        "how is this",
    ]
    alias_index = normalized.find(alias)
    window = normalized[max(0, alias_index - 40): alias_index + len(alias) + 120]
    return any(marker in window for marker in handoff_markers)


def add_trusted_speaker_evidence(
    trusted: Dict[Tuple[int, str], Dict[str, Any]],
    key: Tuple[int, str],
    speaker_id: str,
    reason: str,
    confidence: float,
    text: str = "",
) -> None:
    existing = trusted.get(key)
    if existing and float(existing.get("confidence") or 0) > confidence:
        return
    trusted[key] = {
        "speaker_id": speaker_id,
        "reason": reason,
        "confidence": confidence,
    }
    if text:
        trusted[key]["text"] = clean_segment_text(text)[:240]


def find_next_response_segment(
    segments: List[Dict[str, Any]],
    start_index: int,
    max_gap_seconds: float = 120.0,
) -> Optional[Dict[str, Any]]:
    current = segments[start_index]
    current_chunk = int(current.get("chunk_index") or 0)
    current_local = str(current.get("local_speaker") or current.get("speaker") or "speaker")
    current_end = float(current.get("end") or current.get("start") or 0)
    for candidate in segments[start_index + 1:]:
        if int(candidate.get("chunk_index") or 0) != current_chunk:
            return None
        candidate_start = float(candidate.get("start") or 0)
        if candidate_start - current_end > max_gap_seconds:
            return None
        candidate_local = str(candidate.get("local_speaker") or candidate.get("speaker") or "speaker")
        if candidate_local == current_local:
            continue
        text = clean_segment_text(candidate.get("text", ""))
        if len(text) < 4:
            continue
        return candidate
    return None


def starts_show_as_host(text: str) -> bool:
    normalized = normalize_identity_text(text)
    return any(
        marker in normalized
        for marker in [
            "welcome back",
            "all in podcast",
            "number one podcast",
            "besties are back",
            "everybody s here",
            "original quartet",
        ]
    )


def is_likely_host_moderation_handoff(text: str) -> bool:
    normalized = normalize_identity_text(text)
    return any(
        marker in normalized
        for marker in [
            "you want to add",
            "as we wrap",
            "i ll wrap",
            "ill wrap",
            "wrap up",
        ]
    )


def speaker_by_full_name(roster: List[Dict[str, Any]], full_name: str) -> Optional[Dict[str, Any]]:
    target = normalize_identity_text(full_name)
    for speaker in roster:
        if normalize_identity_text(str(speaker.get("label_full") or "")) == target:
            return speaker
    return None


def build_direct_handoff_evidence(
    segments: List[Dict[str, Any]],
    roster: List[Dict[str, Any]],
) -> Dict[Tuple[int, str], Dict[str, Any]]:
    trusted: Dict[Tuple[int, str], Dict[str, Any]] = {}
    if not roster:
        return trusted

    sorted_segments = sorted(
        segments,
        key=lambda segment: (
            float(segment.get("start") or 0),
            int(segment.get("chunk_index") or 0),
        ),
    )
    first_segment = next((segment for segment in sorted_segments if clean_segment_text(segment.get("text", ""))), None)
    jason = speaker_by_full_name(roster, "Jason Calacanis")
    if first_segment and jason and starts_show_as_host(first_segment.get("text", "")):
        add_trusted_speaker_evidence(
            trusted,
            get_segment_local_key(first_segment),
            str(jason["id"]),
            "show-host-opening",
            0.95,
            str(first_segment.get("text") or ""),
        )

    for index, segment in enumerate(sorted_segments):
        text = clean_segment_text(segment.get("text", ""))
        if not text:
            continue
        for alias, speaker in find_roster_mentions(text, roster):
            if not is_likely_handoff_phrase(text, alias):
                continue
            response = find_next_response_segment(sorted_segments, index)
            if not response:
                continue
            add_trusted_speaker_evidence(
                trusted,
                get_segment_local_key(response),
                str(speaker["id"]),
                "direct-address-response",
                0.98,
                text,
            )
            jason = speaker_by_full_name(roster, "Jason Calacanis")
            if jason and is_likely_host_moderation_handoff(text):
                add_trusted_speaker_evidence(
                    trusted,
                    get_segment_local_key(segment),
                    str(jason["id"]),
                    "host-moderation-handoff",
                    0.92,
                    text,
                )
    return trusted


def text_looks_like_sentence_continuation(previous_text: str, next_text: str) -> bool:
    previous = clean_segment_text(previous_text)
    following = clean_segment_text(next_text)
    if not previous or not following:
        return False
    first_char = following[0]
    if first_char.islower():
        return True
    if previous.endswith(("...", ",", ";", ":", "-", "—")):
        return True
    return False


def build_boundary_continuity_evidence(
    segments: List[Dict[str, Any]],
    trusted_local_speakers: Dict[Tuple[int, str], Dict[str, Any]],
    max_gap_seconds: float = 1.25,
) -> Dict[Tuple[int, str], Dict[str, Any]]:
    evidence: Dict[Tuple[int, str], Dict[str, Any]] = {}
    sorted_segments = sorted(
        segments,
        key=lambda segment: (
            float(segment.get("start") or 0),
            int(segment.get("chunk_index") or 0),
        ),
    )
    for previous, current in zip(sorted_segments, sorted_segments[1:]):
        previous_chunk = int(previous.get("chunk_index") or 0)
        current_chunk = int(current.get("chunk_index") or 0)
        if current_chunk != previous_chunk + 1:
            continue
        previous_key = get_segment_local_key(previous)
        previous_evidence = trusted_local_speakers.get(previous_key)
        if not previous_evidence:
            continue
        previous_end = float(previous.get("end") or previous.get("start") or 0)
        current_start = float(current.get("start") or 0)
        if abs(current_start - previous_end) > max_gap_seconds:
            continue
        if not text_looks_like_sentence_continuation(
            str(previous.get("text") or ""),
            str(current.get("text") or ""),
        ):
            continue
        evidence[get_segment_local_key(current)] = {
            "speaker_id": previous_evidence["speaker_id"],
            "reason": "chunk-boundary-continuation",
            "confidence": min(float(previous_evidence.get("confidence") or 0.9), 0.94),
            "text": (
                clean_segment_text(previous.get("text", ""))
                + " / "
                + clean_segment_text(current.get("text", ""))
            )[:240],
        }
    return evidence


def build_speaker_identity_evidence(
    segments: List[Dict[str, Any]],
    roster: List[Dict[str, Any]],
) -> Dict[str, Any]:
    trusted = build_direct_handoff_evidence(segments, roster)
    boundary = build_boundary_continuity_evidence(segments, trusted)
    for key, item in boundary.items():
        add_trusted_speaker_evidence(
            trusted,
            key,
            str(item["speaker_id"]),
            str(item["reason"]),
            float(item.get("confidence") or 0.9),
            str(item.get("text") or ""),
        )
    return {
        "speakers": [speaker_public_fields(speaker) for speaker in roster],
        "trusted_local_speakers": trusted,
    }


def local_speaker_key_to_json(key: Tuple[int, str]) -> str:
    return f"{int(key[0])}:{key[1]}"


def serialize_speaker_identity_evidence(evidence: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "speakers": evidence.get("speakers", []),
        "trusted_local_speakers": {
            local_speaker_key_to_json(key): value
            for key, value in evidence.get("trusted_local_speakers", {}).items()
        },
    }


def speaker_identity_match_keys(speaker: Dict[str, Any]) -> List[str]:
    keys = [
        str(speaker.get("id") or ""),
        str(speaker.get("label_short") or ""),
        str(speaker.get("label_full") or ""),
        *[str(alias) for alias in speaker.get("aliases", [])],
    ]
    return [normalize_identity_text(key) for key in keys if normalize_identity_text(key)]


def apply_speaker_identity_evidence(
    speaker_mapping: Dict[str, Any],
    evidence: Dict[str, Any],
) -> Dict[str, Any]:
    evidence_speakers = [
        speaker_public_fields(speaker)
        for speaker in evidence.get("speakers", [])
        if speaker.get("id")
    ]
    evidence_by_id = {speaker["id"]: speaker for speaker in evidence_speakers}
    canonical_by_key: Dict[str, str] = {}
    for raw_speaker, public_speaker in zip(evidence.get("speakers", []), evidence_speakers):
        for key in speaker_identity_match_keys(raw_speaker):
            canonical_by_key[key] = public_speaker["id"]
        for key in speaker_identity_match_keys(public_speaker):
            canonical_by_key[key] = public_speaker["id"]

    existing_by_id = {
        str(speaker.get("id")): speaker
        for speaker in speaker_mapping.get("speakers", [])
        if speaker.get("id")
    }
    model_id_to_canonical: Dict[str, str] = {}
    kept_existing_speakers: Dict[str, Dict[str, str]] = {}
    for speaker_id, speaker in existing_by_id.items():
        canonical_id = None
        for key in speaker_identity_match_keys(speaker):
            canonical_id = canonical_by_key.get(key)
            if canonical_id:
                break
        if canonical_id:
            model_id_to_canonical[speaker_id] = canonical_id
        else:
            kept_existing_speakers[speaker_id] = speaker_public_fields(speaker)

    local_speakers: List[Dict[str, Any]] = []
    local_by_key: Dict[Tuple[int, str], Dict[str, Any]] = {}
    for item in speaker_mapping.get("local_speakers", []):
        chunk_index = int(item.get("chunk_index") or 0)
        local_speaker = str(item.get("local_speaker") or "speaker")
        speaker_id = str(item.get("speaker_id") or "")
        speaker_id = model_id_to_canonical.get(speaker_id, speaker_id)
        copied = {
            "chunk_index": chunk_index,
            "local_speaker": local_speaker,
            "speaker_id": speaker_id,
        }
        if item.get("speaker_id_source"):
            copied["speaker_id_source"] = item["speaker_id_source"]
        local_speakers.append(copied)
        local_by_key[(chunk_index, local_speaker)] = copied

    for key, item in evidence.get("trusted_local_speakers", {}).items():
        chunk_index, local_speaker = key
        target = local_by_key.get((int(chunk_index), str(local_speaker)))
        if not target:
            target = {
                "chunk_index": int(chunk_index),
                "local_speaker": str(local_speaker),
                "speaker_id": "",
            }
            local_speakers.append(target)
            local_by_key[(int(chunk_index), str(local_speaker))] = target
        target["speaker_id"] = str(item.get("speaker_id") or target.get("speaker_id") or "")
        target["speaker_id_source"] = str(item.get("reason") or "identity-evidence")

    speakers: List[Dict[str, str]] = list(evidence_by_id.values())
    speakers.extend(
        speaker
        for speaker_id, speaker in kept_existing_speakers.items()
        if speaker_id not in evidence_by_id
    )
    return {
        "speakers": speakers,
        "local_speakers": local_speakers,
    }


def get_local_speaker_mapping_schema(profile_count: int) -> Dict[str, Any]:
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "speakers": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "id": {"type": "string"},
                        "label_short": {"type": "string"},
                        "label_full": {"type": "string"},
                    },
                    "required": ["id", "label_short", "label_full"],
                },
            },
            "local_speakers": {
                "type": "array",
                "minItems": profile_count,
                "maxItems": profile_count,
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "chunk_index": {"type": "integer"},
                        "local_speaker": {"type": "string"},
                        "speaker_id": {"type": "string"},
                    },
                    "required": ["chunk_index", "local_speaker", "speaker_id"],
                },
            },
        },
        "required": ["speakers", "local_speakers"],
    }


def build_local_speaker_mapping_system_prompt() -> str:
    return (
        "You reconcile speaker identities across separately transcribed audio chunks. "
        "The ASR speaker labels are local to each chunk and may reset in every chunk. "
        "For example, local speaker A in chunk 1 may be a different person from local speaker A in chunk 2. "
        "Use the video title, description, timestamps, and sample text to infer stable global speakers. "
        "Map every local chunk speaker to one stable global speaker ID. "
        "If a real name or role can be inferred reliably, use it in the labels; otherwise use Speaker 1, Speaker 2, etc. "
        "Return only JSON that matches the provided schema."
    )


def build_local_speaker_mapping_user_prompt(
    url: str,
    title: str,
    description: str,
    profiles: List[Dict[str, Any]],
    source_language_hint: Optional[str],
    metadata: Optional[Dict[str, Any]] = None,
) -> str:
    lines = [
        f"Video URL: {url}",
        f"Title: {title}",
    ]
    if metadata:
        if metadata.get("channelTitle"):
            lines.append(f"Channel title: {metadata.get('channelTitle')}")
        if metadata.get("tags"):
            lines.append("Tags: " + ", ".join(str(tag) for tag in metadata.get("tags", [])))
    lines.append(f"Description: {description}")
    if source_language_hint:
        lines.append(f"Source language hint: {source_language_hint}")
    lines.append("")
    lines.append("Local speaker profiles:")
    for profile in profiles:
        samples = " / ".join(profile.get("samples", []))
        lines.append(
            f"- chunk {profile['chunk_index']} local speaker {profile['local_speaker']} "
            f"({format_timecode(profile.get('start'))}-{format_timecode(profile.get('end'))}, "
            f"{profile['segment_count']} segments): {samples}"
        )
    return "\n".join(lines)


def assign_global_speakers_for_diarized_segments(
    client: OpenAI,
    model: str,
    url: str,
    title: str,
    description: str,
    segments: List[Dict[str, Any]],
    source_language_hint: Optional[str],
    debug_sink: Optional[List[Dict[str, Any]]] = None,
    metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    profiles = build_local_speaker_profiles(segments)
    if not profiles:
        return {"speakers": [], "local_speakers": []}
    result = call_openai_with_retry(
        client,
        model,
        build_local_speaker_mapping_system_prompt(),
        build_local_speaker_mapping_user_prompt(
            url,
            title,
            description,
            profiles,
            source_language_hint,
            metadata=metadata,
        ),
        schema_name="local_speaker_mapping",
        schema=get_local_speaker_mapping_schema(len(profiles)),
        temperature=0.0,
    )
    if debug_sink is not None:
        debug_sink.append({"profiles": profiles, "result": result})
    return result


def speaker_label_lookup_key(label: str) -> str:
    return re.sub(r"\s+", " ", label.strip().lower())


def resolve_override_speaker_id(
    speakers_by_id: Dict[str, Dict[str, str]],
    override: Dict[str, Any],
) -> str:
    speaker_id = str(override.get("speaker_id") or "").strip()
    if speaker_id:
        return speaker_id

    label = str(
        override.get("speaker_label")
        or override.get("label_short")
        or override.get("label_full")
        or ""
    ).strip()
    if not label:
        raise RuntimeError("Speaker override must include speaker_id or speaker_label")

    label_key = speaker_label_lookup_key(label)
    for existing_id, speaker in speakers_by_id.items():
        labels = [
            existing_id,
            speaker.get("label_short", ""),
            speaker.get("label_full", ""),
        ]
        if any(speaker_label_lookup_key(item) == label_key for item in labels if item):
            return existing_id

    speaker_id = speaker_id_from_label(label)
    speakers_by_id[speaker_id] = {
        "id": speaker_id,
        "label_short": label,
        "label_full": str(override.get("label_full") or label),
    }
    return speaker_id


def apply_speaker_mapping_overrides(
    speaker_mapping: Dict[str, Any],
    overrides: Optional[Dict[str, Any]],
) -> Dict[str, Any]:
    if not overrides:
        return speaker_mapping

    speakers_by_id = {
        str(speaker.get("id")): dict(speaker)
        for speaker in speaker_mapping.get("speakers", [])
        if speaker.get("id")
    }
    for speaker in overrides.get("speakers", []):
        speaker_id = str(speaker.get("id") or "").strip()
        if not speaker_id:
            continue
        speakers_by_id[speaker_id] = {
            "id": speaker_id,
            "label_short": str(speaker.get("label_short") or speaker_id),
            "label_full": str(speaker.get("label_full") or speaker.get("label_short") or speaker_id),
        }

    local_speakers = [
        {
            "chunk_index": int(item.get("chunk_index") or 0),
            "local_speaker": str(item.get("local_speaker") or "speaker"),
            "speaker_id": str(item.get("speaker_id") or ""),
        }
        for item in speaker_mapping.get("local_speakers", [])
    ]
    local_by_key = {
        (item["chunk_index"], item["local_speaker"]): item
        for item in local_speakers
    }

    for override in overrides.get("local_speakers", []):
        chunk_index = int(override.get("chunk_index") or 0)
        local_speaker = str(override.get("local_speaker") or "").strip()
        if not chunk_index or not local_speaker:
            raise RuntimeError("Speaker override must include chunk_index and local_speaker")
        speaker_id = resolve_override_speaker_id(speakers_by_id, override)
        key = (chunk_index, local_speaker)
        if key in local_by_key:
            local_by_key[key]["speaker_id"] = speaker_id
        else:
            item = {
                "chunk_index": chunk_index,
                "local_speaker": local_speaker,
                "speaker_id": speaker_id,
            }
            local_speakers.append(item)
            local_by_key[key] = item

    return {
        "speakers": list(speakers_by_id.values()),
        "local_speakers": local_speakers,
    }


def load_speaker_mapping_overrides(
    video_id: str,
    extra_paths: Optional[List[str]] = None,
) -> Optional[Dict[str, Any]]:
    paths = [get_speaker_mapping_override_path(video_id)]
    if extra_paths:
        paths.extend(extra_paths)

    combined: Dict[str, List[Dict[str, Any]]] = {"speakers": [], "local_speakers": []}
    seen_paths = set()
    found = False
    for path in paths:
        if not path or path in seen_paths or not os.path.exists(path):
            continue
        seen_paths.add(path)
        data = read_json_file(path)
        combined["speakers"].extend(data.get("speakers", []))
        combined["local_speakers"].extend(data.get("local_speakers", []))
        found = True

    return combined if found else None


def get_speaker_mapping_by_local(speaker_mapping: Optional[Dict[str, Any]]) -> Dict[Any, str]:
    if not speaker_mapping:
        return {}
    return {
        (int(item.get("chunk_index") or 0), str(item.get("local_speaker") or "speaker")): str(
            item.get("speaker_id") or ""
        )
        for item in speaker_mapping.get("local_speakers", [])
        if item.get("speaker_id")
    }


def get_segment_local_key(segment: Dict[str, Any]) -> Any:
    chunk_index = int(segment.get("chunk_index") or 0)
    local_speaker = str(segment.get("local_speaker") or segment.get("speaker") or "speaker")
    return chunk_index, local_speaker


def get_baseline_segment_speaker_id(
    segment: Dict[str, Any],
    mapping_by_local: Dict[Any, str],
) -> str:
    chunk_index, local_speaker = get_segment_local_key(segment)
    speaker_id = mapping_by_local.get((chunk_index, local_speaker))
    if speaker_id:
        return speaker_id
    raw_speaker = str(segment.get("speaker") or local_speaker or "speaker")
    return speaker_id_from_label(raw_speaker)


def normalize_embedding(values: List[float]) -> Optional[List[float]]:
    vector = [float(value) for value in values]
    norm = math.sqrt(sum(value * value for value in vector))
    if norm <= 0:
        return None
    return [value / norm for value in vector]


def cosine_similarity(left: List[float], right: List[float]) -> float:
    return sum(a * b for a, b in zip(left, right))


def average_embeddings(embeddings: List[List[float]]) -> Optional[List[float]]:
    if not embeddings:
        return None
    width = len(embeddings[0])
    averaged = [
        sum(embedding[index] for embedding in embeddings) / len(embeddings)
        for index in range(width)
    ]
    return normalize_embedding(averaged)


def robust_embedding_centroid(embeddings: List[List[float]]) -> Optional[List[float]]:
    if len(embeddings) <= 2:
        return average_embeddings(embeddings)
    initial = average_embeddings(embeddings)
    if not initial:
        return None
    ranked = sorted(
        embeddings,
        key=lambda embedding: cosine_similarity(embedding, initial),
        reverse=True,
    )
    keep_count = max(VOICE_RECONCILIATION_MIN_ANCHOR_SEGMENTS, math.ceil(len(ranked) * 0.7))
    return average_embeddings(ranked[:keep_count])


def choose_voice_anchor_groups(
    segments: List[Dict[str, Any]],
    speaker_mapping: Dict[str, Any],
    embeddings: Dict[int, List[float]],
) -> Dict[str, Dict[str, Any]]:
    mapping_by_local = get_speaker_mapping_by_local(speaker_mapping)
    groups: Dict[Any, Dict[str, Any]] = {}
    for index, segment in enumerate(segments):
        embedding = embeddings.get(index)
        if embedding is None:
            continue
        speaker_id = get_baseline_segment_speaker_id(segment, mapping_by_local)
        chunk_index, local_speaker = get_segment_local_key(segment)
        key = (speaker_id, chunk_index, local_speaker)
        start = float(segment.get("start") or 0)
        end = float(segment.get("end") or start)
        group = groups.setdefault(
            key,
            {
                "speaker_id": speaker_id,
                "chunk_index": chunk_index,
                "local_speaker": local_speaker,
                "start": start,
                "end": end,
                "total_duration": 0.0,
                "embeddings": [],
                "segment_indexes": [],
            },
        )
        group["start"] = min(float(group["start"]), start)
        group["end"] = max(float(group["end"]), end)
        group["total_duration"] += max(0.0, end - start)
        group["embeddings"].append(embedding)
        group["segment_indexes"].append(index)

    anchors: Dict[str, Dict[str, Any]] = {}
    speaker_ids = [
        str(speaker.get("id"))
        for speaker in speaker_mapping.get("speakers", [])
        if speaker.get("id")
    ]
    for speaker_id in speaker_ids:
        candidates = [
            group
            for group in groups.values()
            if group["speaker_id"] == speaker_id
        ]
        candidates.sort(
            key=lambda group: (
                int(group["chunk_index"]),
                float(group["start"]),
                -float(group["total_duration"]),
            )
        )
        qualified = [
            group
            for group in candidates
            if len(group["embeddings"]) >= VOICE_RECONCILIATION_MIN_ANCHOR_SEGMENTS
            and float(group["total_duration"]) >= VOICE_RECONCILIATION_MIN_ANCHOR_SECONDS
        ]
        chosen = qualified[0] if qualified else (candidates[0] if candidates else None)
        if not chosen:
            continue
        embeddings_for_centroid = chosen["embeddings"][:VOICE_RECONCILIATION_MAX_ANCHOR_SEGMENTS]
        centroid = robust_embedding_centroid(embeddings_for_centroid)
        if not centroid:
            continue
        anchors[speaker_id] = {
            "speaker_id": speaker_id,
            "chunk_index": chosen["chunk_index"],
            "local_speaker": chosen["local_speaker"],
            "start": round(float(chosen["start"]), 3),
            "end": round(float(chosen["end"]), 3),
            "segment_count": len(chosen["embeddings"]),
            "total_duration": round(float(chosen["total_duration"]), 3),
            "centroid": centroid,
        }
    return anchors


def reconcile_segment_speakers_with_voice_embeddings(
    segments: List[Dict[str, Any]],
    speaker_mapping: Dict[str, Any],
    segment_embeddings: Dict[int, List[float]],
    min_similarity: float = VOICE_RECONCILIATION_MIN_SIMILARITY,
    min_margin: float = VOICE_RECONCILIATION_MIN_MARGIN,
    neighbor_gap_seconds: float = VOICE_RECONCILIATION_NEIGHBOR_GAP_SECONDS,
) -> Any:
    mapping_by_local = get_speaker_mapping_by_local(speaker_mapping)
    normalized_embeddings: Dict[int, List[float]] = {}
    for index, embedding in segment_embeddings.items():
        normalized = normalize_embedding(embedding)
        if normalized is not None:
            normalized_embeddings[int(index)] = normalized

    anchors = choose_voice_anchor_groups(segments, speaker_mapping, normalized_embeddings)
    speakers_by_id = {
        str(speaker.get("id")): speaker
        for speaker in speaker_mapping.get("speakers", [])
        if speaker.get("id")
    }
    debug: Dict[str, Any] = {
        "status": "ok" if len(anchors) >= 2 else "skipped",
        "speaker_anchor_groups": [
            {
                key: value
                for key, value in anchor.items()
                if key != "centroid"
            }
            for anchor in anchors.values()
        ],
        "segment_count": len(segments),
        "embedded_segment_count": len(normalized_embeddings),
        "voice_assigned_count": 0,
        "voice_changed_count": 0,
        "local_majority_assigned_count": 0,
        "neighbor_assigned_count": 0,
        "changed_examples": [],
    }
    resolved = [dict(segment) for segment in segments]
    if len(anchors) < 2:
        for segment in resolved:
            speaker_id = get_baseline_segment_speaker_id(segment, mapping_by_local)
            segment.setdefault("speaker_id", speaker_id)
            segment.setdefault("speaker_id_source", "local_mapping")
        return resolved, debug

    anchor_items = list(anchors.items())
    for index, segment in enumerate(resolved):
        baseline_speaker_id = get_baseline_segment_speaker_id(segment, mapping_by_local)
        segment["speaker_id"] = baseline_speaker_id
        segment["speaker_id_source"] = "local_mapping"
        embedding = normalized_embeddings.get(index)
        if embedding is None:
            continue

        scored = sorted(
            (
                (speaker_id, cosine_similarity(embedding, anchor["centroid"]))
                for speaker_id, anchor in anchor_items
            ),
            key=lambda item: item[1],
            reverse=True,
        )
        best_speaker_id, best_score = scored[0]
        second_score = scored[1][1] if len(scored) > 1 else -1.0
        margin = best_score - second_score
        segment["voice_speaker_id"] = best_speaker_id
        segment["voice_similarity"] = round(best_score, 4)
        segment["voice_similarity_margin"] = round(margin, 4)
        if best_score < min_similarity or margin < min_margin:
            continue

        segment["speaker_id"] = best_speaker_id
        segment["speaker_id_source"] = "voice"
        debug["voice_assigned_count"] += 1
        if best_speaker_id != baseline_speaker_id:
            debug["voice_changed_count"] += 1
            if len(debug["changed_examples"]) < 20:
                debug["changed_examples"].append(
                    {
                        "index": index + 1,
                        "start": round(float(segment.get("start") or 0), 3),
                        "chunk_index": int(segment.get("chunk_index") or 0),
                        "local_speaker": str(segment.get("local_speaker") or segment.get("speaker") or "speaker"),
                        "baseline_speaker_id": baseline_speaker_id,
                        "voice_speaker_id": best_speaker_id,
                        "voice_similarity": round(best_score, 4),
                        "voice_similarity_margin": round(margin, 4),
                        "text": clean_segment_text(segment.get("text", ""))[:180],
                    }
                )

    local_voice_counts: Dict[Any, Dict[str, int]] = {}
    for segment in resolved:
        if segment.get("speaker_id_source") != "voice":
            continue
        local_key = get_segment_local_key(segment)
        speaker_id = str(segment.get("speaker_id") or "")
        if not speaker_id:
            continue
        local_voice_counts.setdefault(local_key, {})
        local_voice_counts[local_key][speaker_id] = local_voice_counts[local_key].get(speaker_id, 0) + 1

    for segment in resolved:
        if segment.get("speaker_id_source") != "local_mapping":
            continue
        counts = local_voice_counts.get(get_segment_local_key(segment), {})
        if not counts:
            continue
        top_speaker_id, top_count = max(counts.items(), key=lambda item: item[1])
        total_count = sum(counts.values())
        if (
            top_count < VOICE_RECONCILIATION_LOCAL_MAJORITY_MIN_SEGMENTS
            or top_count / total_count < VOICE_RECONCILIATION_LOCAL_MAJORITY_SHARE
        ):
            continue
        if top_speaker_id == segment.get("speaker_id"):
            continue
        segment["speaker_id"] = top_speaker_id
        segment["speaker_id_source"] = "voice_local_majority"
        debug["local_majority_assigned_count"] += 1

    propagating_voice_sources = {"voice", "voice_neighbor", "voice_local_majority"}
    changed = True
    while changed:
        changed = False
        for index, segment in enumerate(resolved):
            if segment.get("speaker_id_source") != "local_mapping":
                continue
            current_key = get_segment_local_key(segment)
            current_start = float(segment.get("start") or 0)
            previous = resolved[index - 1] if index > 0 else None
            if previous and previous.get("speaker_id_source") in propagating_voice_sources:
                previous_key = get_segment_local_key(previous)
                previous_end = float(previous.get("end") or previous.get("start") or 0)
                if previous_key == current_key and current_start - previous_end <= neighbor_gap_seconds:
                    segment["speaker_id"] = str(previous["speaker_id"])
                    segment["speaker_id_source"] = "voice_neighbor"
                    debug["neighbor_assigned_count"] += 1
                    changed = True
                    continue

            next_segment = resolved[index + 1] if index + 1 < len(resolved) else None
            if next_segment and next_segment.get("speaker_id_source") in propagating_voice_sources:
                next_key = get_segment_local_key(next_segment)
                current_end = float(segment.get("end") or segment.get("start") or 0)
                next_start = float(next_segment.get("start") or 0)
                if next_key == current_key and next_start - current_end <= neighbor_gap_seconds:
                    segment["speaker_id"] = str(next_segment["speaker_id"])
                    segment["speaker_id_source"] = "voice_neighbor"
                    debug["neighbor_assigned_count"] += 1
                    changed = True

    debug["speaker_count"] = len(speakers_by_id)
    return resolved, debug


def segment_duration_or_unit(segment: Dict[str, Any]) -> float:
    start = segment.get("start")
    end = segment.get("end")
    try:
        if start is not None and end is not None:
            duration = float(end) - float(start)
            if duration > 0:
                return duration
    except (TypeError, ValueError):
        pass
    return 1.0


def is_role_like_speaker(speaker: Dict[str, str]) -> bool:
    label = normalize_identity_text(
        " ".join(
            [
                str(speaker.get("id") or ""),
                str(speaker.get("label_short") or ""),
                str(speaker.get("label_full") or ""),
            ]
        )
    )
    return any(marker in label for marker in ROLE_SPEAKER_LABEL_MARKERS)


def collapse_role_speaker_identities(
    segments: List[Dict[str, Any]],
    speakers: List[Dict[str, str]],
) -> Tuple[List[Dict[str, Any]], List[Dict[str, str]], Dict[str, Any]]:
    role_speaker_ids = {
        str(speaker.get("id"))
        for speaker in speakers
        if speaker.get("id") and is_role_like_speaker(speaker)
    }
    real_speaker_ids = {
        str(speaker.get("id"))
        for speaker in speakers
        if speaker.get("id") and str(speaker.get("id")) not in role_speaker_ids
    }
    debug: Dict[str, Any] = {
        "role_speaker_ids": sorted(role_speaker_ids),
        "merged_role_speakers": {},
        "candidate_scores": {},
        "changed_segment_count": 0,
    }
    if not role_speaker_ids or not real_speaker_ids:
        return [dict(segment) for segment in segments], list(speakers), debug

    groups: Dict[Any, List[Dict[str, Any]]] = {}
    for segment in segments:
        groups.setdefault(get_segment_local_key(segment), []).append(segment)

    candidate_scores: Dict[str, Dict[str, float]] = {role_id: {} for role_id in role_speaker_ids}
    for group_segments in groups.values():
        group_role_ids = {
            str(segment.get("speaker_id") or "")
            for segment in group_segments
            if str(segment.get("speaker_id") or "") in role_speaker_ids
        }
        if not group_role_ids:
            continue
        for segment in group_segments:
            speaker_id = str(segment.get("speaker_id") or "")
            if speaker_id not in real_speaker_ids:
                continue
            if segment.get("speaker_id_source") not in VOICE_DERIVED_SPEAKER_SOURCES:
                continue
            weight = segment_duration_or_unit(segment)
            for role_id in group_role_ids:
                candidate_scores[role_id][speaker_id] = candidate_scores[role_id].get(speaker_id, 0.0) + weight

    merge_by_role: Dict[str, str] = {}
    for role_id, scores in candidate_scores.items():
        if not scores:
            continue
        target_id, score = max(scores.items(), key=lambda item: item[1])
        if score <= 0:
            continue
        merge_by_role[role_id] = target_id

    debug["candidate_scores"] = {
        role_id: {speaker_id: round(score, 3) for speaker_id, score in scores.items()}
        for role_id, scores in candidate_scores.items()
        if scores
    }
    debug["merged_role_speakers"] = dict(merge_by_role)
    if not merge_by_role:
        return [dict(segment) for segment in segments], list(speakers), debug

    merged_segments: List[Dict[str, Any]] = []
    changed_segment_count = 0
    for segment in segments:
        copied = dict(segment)
        speaker_id = str(copied.get("speaker_id") or "")
        target_id = merge_by_role.get(speaker_id)
        if target_id:
            copied["speaker_id_before_role_merge"] = speaker_id
            copied["speaker_id"] = target_id
            copied["speaker_id_source_before_role_merge"] = copied.get("speaker_id_source")
            copied["speaker_id_source"] = "role_identity_merge"
            changed_segment_count += 1
        merged_segments.append(copied)

    merged_speakers = [
        speaker
        for speaker in speakers
        if str(speaker.get("id") or "") not in merge_by_role
    ]
    debug["changed_segment_count"] = changed_segment_count
    return merged_segments, merged_speakers, debug


def get_voice_reconciliation_enabled() -> bool:
    raw = os.getenv("YTRANSLATE_VOICE_SPEAKER_RECONCILIATION", "1").strip().lower()
    return raw not in {"0", "false", "no", "off"}


def transcode_audio_for_voice_reconciliation(
    source_audio_path: str,
    video_id: str,
    log: Callable[[str], None],
) -> str:
    voice_dir = os.path.join(get_video_cache_dir(video_id), "voice-speaker-reconciliation")
    os.makedirs(voice_dir, exist_ok=True)
    output_path = os.path.join(voice_dir, "source-16k.wav")
    if os.path.exists(output_path):
        return output_path

    command = [
        get_ffmpeg_executable(),
        "-hide_banner",
        "-y",
        "-i",
        source_audio_path,
        "-vn",
        "-ac",
        "1",
        "-ar",
        "16000",
        output_path,
    ]
    log("Preparing audio for voice speaker reconciliation...")
    subprocess.run(command, check=True)
    return output_path


def build_voice_segment_embeddings(
    wav_path: str,
    segments: List[Dict[str, Any]],
    log: Callable[[str], None],
    min_segment_seconds: float = VOICE_RECONCILIATION_MIN_SEGMENT_SECONDS,
    max_segment_seconds: float = VOICE_RECONCILIATION_MAX_SEGMENT_SECONDS,
) -> Dict[int, List[float]]:
    try:
        import soundfile as sf
        from resemblyzer import VoiceEncoder
    except ImportError as exc:
        raise RuntimeError(
            "Voice speaker reconciliation requires resemblyzer and soundfile."
        ) from exc

    audio, sample_rate = sf.read(wav_path, dtype="float32")
    if getattr(audio, "ndim", 1) > 1:
        audio = audio.mean(axis=1)
    if int(sample_rate) != 16000:
        raise RuntimeError(f"Expected 16000 Hz voice reconciliation audio, got {sample_rate}")

    encoder = VoiceEncoder()
    embeddings: Dict[int, List[float]] = {}
    audio_length = len(audio)
    for index, segment in enumerate(segments):
        start = float(segment.get("start") or 0)
        end = float(segment.get("end") or start)
        duration = end - start
        if duration < min_segment_seconds:
            continue
        if duration > max_segment_seconds:
            midpoint = (start + end) / 2
            start = midpoint - max_segment_seconds / 2
            end = midpoint + max_segment_seconds / 2
        start_sample = max(0, int(start * sample_rate))
        end_sample = min(audio_length, int(end * sample_rate))
        if end_sample - start_sample < int(min_segment_seconds * sample_rate):
            continue
        embedding = encoder.embed_utterance(audio[start_sample:end_sample])
        embeddings[index] = [float(value) for value in embedding]

    log(f"Voice speaker reconciliation embedded {len(embeddings)} transcript segments.")
    return embeddings


def reconcile_diarized_segments_with_voice(
    url: str,
    video_id: str,
    segments: List[Dict[str, Any]],
    speaker_mapping: Dict[str, Any],
    log: Callable[[str], None],
) -> Any:
    if not get_voice_reconciliation_enabled():
        return [dict(segment) for segment in segments], {"status": "disabled"}
    if len(speaker_mapping.get("speakers", [])) < 2 or not segments:
        return [dict(segment) for segment in segments], {"status": "skipped"}

    try:
        source_audio_path = download_youtube_audio(url, video_id, log)
        wav_path = transcode_audio_for_voice_reconciliation(source_audio_path, video_id, log)
        embeddings = build_voice_segment_embeddings(wav_path, segments, log)
        resolved, debug = reconcile_segment_speakers_with_voice_embeddings(
            segments,
            speaker_mapping,
            embeddings,
        )
        if debug.get("status") == "ok":
            log(
                "Voice speaker reconciliation changed "
                f"{debug.get('voice_changed_count', 0)} segments "
                f"and filled {debug.get('local_majority_assigned_count', 0)} local-majority "
                f"plus {debug.get('neighbor_assigned_count', 0)} short-neighbor segments."
            )
        return resolved, debug
    except Exception as exc:
        log(f"Voice speaker reconciliation skipped ({exc}).")
        resolved = [dict(segment) for segment in segments]
        mapping_by_local = get_speaker_mapping_by_local(speaker_mapping)
        for segment in resolved:
            segment.setdefault("speaker_id", get_baseline_segment_speaker_id(segment, mapping_by_local))
            segment.setdefault("speaker_id_source", "local_mapping")
        return resolved, {"status": "skipped", "error": str(exc)}


def get_speaker_identity_linker_enabled() -> bool:
    raw = os.getenv("YTRANSLATE_SPEAKER_IDENTITY_LINKER", "1").strip().lower()
    return raw not in {"0", "false", "no", "off"}


def speaker_identity_boundary_schema(segment_count: int) -> Dict[str, Any]:
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "boundaries": {
                "type": "array",
                "minItems": segment_count,
                "maxItems": segment_count,
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "segment_id": {"type": "integer"},
                        "boundary_before": {
                            "type": "string",
                            "enum": ["same", "uncertain", "change"],
                        },
                    },
                    "required": ["segment_id", "boundary_before"],
                },
            }
        },
        "required": ["boundaries"],
    }


def speaker_identity_boundary_system_prompt() -> str:
    return (
        "Detect anonymous speaker-turn boundaries in an ASR transcript. Do not identify or name speakers. "
        "For every TARGET segment, classify whether the same real person continues from the immediately "
        "preceding transcript segment: 'same' means confidently the same person, 'change' means confidently "
        "a different person starts, and 'uncertain' means the evidence is insufficient. Local diarizer labels "
        "are weak hints and can both merge two people and split one uninterrupted speaker. Use grammar, sentence "
        "continuity, overlapping timestamps, direct address, question-answer flow, interruptions, and discourse. "
        "Use 'same' or 'change' only when confident because they become clustering constraints; otherwise choose "
        "'uncertain'. Return each TARGET segment exactly once with its exact segment_id."
    )


def speaker_identity_boundary_batches(
    segment_count: int,
    batch_size: int = SPEAKER_IDENTITY_LINKER_BATCH_SEGMENTS,
    context_size: int = SPEAKER_IDENTITY_LINKER_CONTEXT_SEGMENTS,
) -> List[Tuple[range, range]]:
    batches: List[Tuple[range, range]] = []
    for start in range(0, segment_count, batch_size):
        stop = min(segment_count, start + batch_size)
        context_start = max(0, start - context_size)
        context_stop = min(segment_count, stop + context_size)
        batches.append((range(start, stop), range(context_start, context_stop)))
    return batches


def speaker_identity_boundary_user_prompt(
    title: str,
    context: str,
    segments: List[Dict[str, Any]],
    target_ids: range,
    context_ids: range,
) -> str:
    target_set = set(target_ids)
    lines = [
        f"Title: {title}",
        f"Context: {context}",
        "",
        "Transcript. Classify TARGET lines only; BEFORE and AFTER are context:",
    ]
    for segment_id in context_ids:
        segment = segments[segment_id]
        if segment_id in target_set:
            marker = "TARGET"
        elif segment_id < target_ids.start:
            marker = "BEFORE"
        else:
            marker = "AFTER"
        chunk_index, local_speaker = get_segment_local_key(segment)
        start = float(segment.get("start") or 0)
        end = float(segment.get("end") or start)
        lines.append(
            f"{marker} {segment_id} | {start:.3f}-{end:.3f} | "
            f"local={chunk_index}/{local_speaker} | "
            f"text={clean_segment_text(str(segment.get('text') or ''))}"
        )
    return "\n".join(lines)


def validate_speaker_identity_boundaries(
    assignments: List[Dict[str, Any]],
    target_ids: range,
) -> Dict[int, str]:
    expected = set(target_ids)
    result: Dict[int, str] = {}
    for assignment in assignments:
        segment_id = int(assignment.get("segment_id"))
        boundary = str(assignment.get("boundary_before") or "")
        if segment_id in result:
            raise RuntimeError(f"Duplicate boundary result for segment {segment_id}")
        if boundary not in {"same", "uncertain", "change"}:
            raise RuntimeError(f"Invalid boundary value {boundary!r} for segment {segment_id}")
        result[segment_id] = boundary
    if set(result) != expected:
        missing = sorted(expected - set(result))
        extra = sorted(set(result) - expected)
        raise RuntimeError(
            f"Boundary IDs mismatch; missing={missing[:10]}, extra={extra[:10]}"
        )
    return result


def infer_speaker_identity_boundaries(
    client: OpenAI,
    video_id: str,
    title: str,
    context: str,
    segments: List[Dict[str, Any]],
    log: Callable[[str], None],
) -> Tuple[List[str], Dict[str, Any]]:
    if not segments:
        return [], {"status": "skipped", "reason": "no-segments"}

    boundaries = ["uncertain"] * len(segments)
    boundaries[0] = "change"
    batches = speaker_identity_boundary_batches(len(segments))
    batch_dir = os.path.join(
        get_video_cache_dir(video_id),
        "speaker-identity-linker",
        "boundary-batches",
        SPEAKER_IDENTITY_LINKER_REASONING_EFFORT,
    )
    os.makedirs(batch_dir, exist_ok=True)
    cache_hits = 0
    requests = 0
    for batch_number, (target_ids, context_ids) in enumerate(batches, 1):
        schema = speaker_identity_boundary_schema(len(target_ids))
        system_prompt = speaker_identity_boundary_system_prompt()
        user_prompt = speaker_identity_boundary_user_prompt(
            title,
            context,
            segments,
            target_ids,
            context_ids,
        )
        digest_payload = {
            "schema_version": SPEAKER_IDENTITY_LINKER_CACHE_SCHEMA_VERSION,
            "model": SPEAKER_IDENTITY_LINKER_MODEL,
            "reasoning_effort": SPEAKER_IDENTITY_LINKER_REASONING_EFFORT,
            "schema": schema,
            "system_prompt": system_prompt,
            "user_prompt": user_prompt,
        }
        input_digest = hashlib.sha256(
            json.dumps(
                digest_payload,
                ensure_ascii=False,
                sort_keys=True,
                separators=(",", ":"),
            ).encode("utf-8")
        ).hexdigest()
        batch_path = os.path.join(batch_dir, f"batch-{batch_number:02d}.json")
        cached: Optional[Dict[str, Any]] = None
        if os.path.exists(batch_path):
            candidate = read_json_file(batch_path)
            if (
                candidate.get("schema_version")
                == SPEAKER_IDENTITY_LINKER_CACHE_SCHEMA_VERSION
                and candidate.get("input_digest") == input_digest
            ):
                cached = candidate
        if cached is not None:
            result = {"boundaries": cached.get("boundaries", [])}
            cache_hits += 1
        else:
            log(
                f"Speaker identity boundary batch {batch_number}/{len(batches)} "
                f"({len(target_ids)} segments)..."
            )
            result = call_openai_with_retry(
                client,
                SPEAKER_IDENTITY_LINKER_MODEL,
                system_prompt,
                user_prompt,
                schema_name=f"speaker_identity_boundaries_{batch_number}",
                schema=schema,
                reasoning_effort=SPEAKER_IDENTITY_LINKER_REASONING_EFFORT,
            )
            requests += 1
            write_json_file(
                batch_path,
                {
                    "schema_version": SPEAKER_IDENTITY_LINKER_CACHE_SCHEMA_VERSION,
                    "input_digest": input_digest,
                    "boundaries": result.get("boundaries", []),
                },
            )
        validated = validate_speaker_identity_boundaries(
            list(result.get("boundaries", [])),
            target_ids,
        )
        for segment_id, boundary in validated.items():
            boundaries[segment_id] = boundary
    boundaries[0] = "change"
    counts = {value: boundaries.count(value) for value in ("same", "uncertain", "change")}
    return boundaries, {
        "status": "ok",
        "model": SPEAKER_IDENTITY_LINKER_MODEL,
        "reasoning_effort": SPEAKER_IDENTITY_LINKER_REASONING_EFFORT,
        "batch_count": len(batches),
        "cache_hits": cache_hits,
        "request_count": requests,
        "boundary_counts": counts,
    }


def attach_speaker_labels_to_segments(
    segments: List[Dict[str, Any]],
    speakers: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    labels_by_id = {
        str(speaker.get("id") or ""): str(
            speaker.get("label_full")
            or speaker.get("label_short")
            or speaker.get("id")
            or ""
        )
        for speaker in speakers
        if speaker.get("id")
    }
    resolved: List[Dict[str, Any]] = []
    for segment in segments:
        copied = dict(segment)
        speaker_id = str(copied.get("speaker_id") or "")
        if speaker_id in labels_by_id:
            copied["speaker_label"] = labels_by_id[speaker_id]
        resolved.append(copied)
    return resolved


def merge_linked_speakers(
    speakers: List[Dict[str, Any]],
    segments: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    merged = {
        str(speaker.get("id") or ""): dict(speaker)
        for speaker in speakers
        if speaker.get("id")
    }
    for segment in segments:
        speaker_id = str(segment.get("speaker_id") or "").strip()
        speaker_label = str(segment.get("speaker_label") or "").strip()
        if not speaker_id or not speaker_label:
            continue
        merged.setdefault(
            speaker_id,
            {
                "id": speaker_id,
                "label_short": speaker_label,
                "label_full": speaker_label,
            },
        )
    return list(merged.values())


def normalize_linked_speaker_ids(
    segments: List[Dict[str, Any]],
    speakers: List[Dict[str, Any]],
) -> List[Dict[str, Any]]:
    try:
        from speaker_identity_linker import canonical_speaker_name
    except ImportError:
        return [dict(segment) for segment in segments]

    ids_by_identity: Dict[str, str] = {}
    for speaker in speakers:
        speaker_id = str(speaker.get("id") or "").strip()
        label = str(
            speaker.get("label_full")
            or speaker.get("label_short")
            or speaker_id
        )
        if speaker_id:
            ids_by_identity.setdefault(canonical_speaker_name(label), speaker_id)

    normalized: List[Dict[str, Any]] = []
    for segment in segments:
        copied = dict(segment)
        label = str(copied.get("speaker_label") or copied.get("speaker_id") or "")
        target_id = ids_by_identity.get(canonical_speaker_name(label))
        current_id = str(copied.get("speaker_id") or "")
        if target_id and target_id != current_id:
            copied["speaker_id_before_identity_normalization"] = current_id
            copied["speaker_id"] = target_id
        normalized.append(copied)
    return normalized


def run_speaker_identity_linker(
    client: OpenAI,
    url: str,
    video_id: str,
    metadata: Dict[str, Any],
    segments: List[Dict[str, Any]],
    speakers: List[Dict[str, Any]],
    log: Callable[[str], None],
) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]], Dict[str, Any]]:
    baseline = attach_speaker_labels_to_segments(segments, speakers)
    if not get_speaker_identity_linker_enabled():
        return baseline, list(speakers), {"status": "disabled"}
    if not infer_known_speaker_roster(metadata):
        return baseline, list(speakers), {
            "status": "skipped",
            "reason": "unsupported-show",
        }
    if len(speakers) < 2 or not baseline:
        return baseline, list(speakers), {
            "status": "skipped",
            "reason": "insufficient-speakers-or-segments",
        }

    try:
        import soundfile as sf
        from resemblyzer import VoiceEncoder

        import speaker_identity_linker as linker
        from all_in_speaker_reference_bank import ALL_IN_REFERENCE_BANK

        source_audio_path = download_youtube_audio(url, video_id, log)
        wav_path = transcode_audio_for_voice_reconciliation(
            source_audio_path,
            video_id,
            log,
        )
        audio, sample_rate = sf.read(wav_path, dtype="float32")
        if getattr(audio, "ndim", 1) > 1:
            audio = audio.mean(axis=1)
        if int(sample_rate) != 16000:
            raise RuntimeError(f"Expected 16000 Hz linker audio, got {sample_rate}")

        encoder = VoiceEncoder()
        references, reference_debug = linker.build_episode_reference_centroids(
            baseline,
            audio,
            int(sample_rate),
            encoder,
            dict(ALL_IN_REFERENCE_BANK.get("speakers", {})),
            active_speaker_labels=[
                str(
                    speaker.get("label_full")
                    or speaker.get("label_short")
                    or speaker.get("id")
                    or ""
                )
                for speaker in speakers
                if linker.canonical_speaker_name(
                    speaker.get("label_full")
                    or speaker.get("label_short")
                    or speaker.get("id")
                )
                not in {"jason", "chamath", "sacks", "friedberg"}
            ],
        )
        if len(references) < 2:
            raise RuntimeError("Fewer than two active speaker voice references")

        active_labels = ", ".join(
            str(speaker.get("label_full") or speaker.get("label_short") or "")
            for speaker in speakers
            if speaker.get("id")
        )
        context = (
            f"Channel: {metadata.get('channelTitle') or 'All-In Podcast'}. "
            f"Candidate panel identities from prior attribution: {active_labels}."
        )
        boundaries, boundary_debug = infer_speaker_identity_boundaries(
            client,
            video_id,
            str(metadata.get("title") or "Untitled"),
            context,
            baseline,
            log,
        )
        resolved, linker_debug = linker.link_speaker_identities(
            baseline,
            boundaries,
            audio,
            int(sample_rate),
            encoder,
            references,
            log=log,
        )
        resolved = normalize_linked_speaker_ids(resolved, speakers)
        merged_speakers = merge_linked_speakers(speakers, resolved)
        changed_count = sum(
            str(before.get("speaker_id") or "") != str(after.get("speaker_id") or "")
            for before, after in zip(baseline, resolved)
        )
        debug = {
            "status": "ok",
            "scope": "recognized-all-in",
            "changed_segment_count": changed_count,
            "reference_bank_schema_version": ALL_IN_REFERENCE_BANK.get("schema_version"),
            "references": reference_debug,
            "boundaries": boundary_debug,
            "linker": linker_debug,
        }
        log(f"Speaker identity linker changed {changed_count} segments.")
        return resolved, merged_speakers, debug
    except Exception as exc:
        log(f"Speaker identity linker skipped ({exc}).")
        return baseline, list(speakers), {"status": "skipped", "error": str(exc)}


def attributed_turns_from_diarized_segments(
    segments: List[Dict[str, Any]],
    speaker_mapping: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    speakers_by_id: Dict[str, Dict[str, str]] = {}
    turns: List[Dict[str, str]] = []
    mapping_by_local = {}
    if speaker_mapping:
        for speaker in speaker_mapping.get("speakers", []):
            speakers_by_id[speaker["id"]] = speaker
        mapping_by_local = get_speaker_mapping_by_local(speaker_mapping)

    for segment in segments:
        raw_speaker = str(segment.get("speaker") or "speaker")
        speaker_label = raw_speaker if raw_speaker.lower().startswith("speaker") else f"Speaker {raw_speaker}"
        local_speaker = str(segment.get("local_speaker") or raw_speaker)
        chunk_index = int(segment.get("chunk_index") or 0)
        speaker_id = str(segment.get("speaker_id") or "").strip()
        if speaker_id and speaker_id not in speakers_by_id:
            speakers_by_id[speaker_id] = {
                "id": speaker_id,
                "label_short": speaker_id,
                "label_full": speaker_id,
            }
        if not speaker_id:
            speaker_id = mapping_by_local.get((chunk_index, local_speaker))
        if not speaker_id:
            speaker_id = speaker_id_from_label(raw_speaker)
            speakers_by_id.setdefault(
                speaker_id,
                {
                    "id": speaker_id,
                    "label_short": speaker_label,
                    "label_full": speaker_label,
                },
            )
        text = clean_segment_text(segment.get("text", ""))
        if not text:
            continue
        if turns and turns[-1].get("speaker_id") == speaker_id:
            turns[-1]["text_source"] = (turns[-1].get("text_source", "") + " " + text).strip()
        else:
            turns.append({"speaker_id": speaker_id, "text_source": text})

    return {
        "speakers": list(speakers_by_id.values()),
        "turns": turns,
    }


def aliases_for_speaker_label(label_short: str, label_full: str) -> List[str]:
    labels = [label_short, label_full]
    normalized_full = normalize_identity_text(label_full)
    normalized_short = normalize_identity_text(label_short)
    if "friedberg" in normalized_full or "friedberg" in normalized_short:
        labels.extend(["Friedberg", "Freeberg", "Freiberg"])
    if "sacks" in normalized_full or "sacks" in normalized_short:
        labels.extend(["Sacks", "Sachs", "Zach"])
    if "chamath" in normalized_full or "chamath" in normalized_short:
        labels.extend(["Chamath", "Chumath", "Jamath"])
    if "jason" in normalized_full or "jason" in normalized_short:
        labels.extend(["Jason", "J-Cal", "JCal", "J Cal"])
    seen = set()
    aliases = []
    for label in labels:
        normalized = normalize_identity_text(label)
        if normalized and normalized not in seen:
            seen.add(normalized)
            aliases.append(normalized)
    return aliases


def text_contains_self_reference_contradiction(text: str, alias: str) -> bool:
    normalized = normalize_identity_text(text)
    if not alias or f" {alias} " not in f" {normalized} ":
        return False
    patterns = [
        f"what {alias} said",
        f"{alias} said",
        f"{alias} says",
        f"to {alias} s point",
        f"{alias} s point",
        f"{alias} is right",
        f"{alias} was right",
        f"{alias} you",
        f"{alias} what",
    ]
    return any(pattern in normalized for pattern in patterns)


def find_speaker_identity_contradictions(
    speakers: List[Dict[str, str]],
    turns: List[Dict[str, str]],
) -> List[Dict[str, Any]]:
    speakers_by_id = {
        str(speaker.get("id")): speaker
        for speaker in speakers
        if speaker.get("id")
    }
    issues: List[Dict[str, Any]] = []
    for index, turn in enumerate(turns, 1):
        speaker_id = str(turn.get("speaker_id") or "")
        speaker = speakers_by_id.get(speaker_id)
        if not speaker:
            continue
        text = clean_segment_text(turn.get("text_source") or turn.get("text_translated") or "")
        for alias in aliases_for_speaker_label(
            str(speaker.get("label_short") or ""),
            str(speaker.get("label_full") or ""),
        ):
            if not text_contains_self_reference_contradiction(text, alias):
                continue
            issues.append(
                {
                    "turn_index": index,
                    "speaker_id": speaker_id,
                    "speaker_label": speaker.get("label_short") or speaker_id,
                    "matched_alias": alias,
                    "text": text[:240],
                }
            )
            break
    return issues


def transcribe_youtube_audio_with_openai(
    url: str,
    video_id: str,
    openai_key: str,
    log: Callable[[str], None],
) -> Dict[str, Any]:
    asr_model = os.getenv("OPENAI_ASR_MODEL", OPENAI_ASR_MODEL)
    chunk_seconds = get_asr_chunk_seconds()
    jobs = get_asr_jobs()
    timeout_seconds = get_asr_timeout_seconds()
    max_passes = get_asr_max_passes()
    retry_pass_delay_seconds = get_asr_retry_pass_delay_seconds()
    cache_dir = get_video_cache_dir(video_id)
    result_path = os.path.join(cache_dir, f"openai-asr-{asr_model}-{chunk_seconds}s.json")
    if os.path.exists(result_path):
        return read_json_file(result_path)

    audio_path = download_youtube_audio(url, video_id, log)
    chunks = transcode_and_chunk_audio(audio_path, video_id, chunk_seconds, log)
    offsets = build_chunk_offsets(chunks)
    log(f"Prepared {len(chunks)} OpenAI ASR chunks ({chunk_seconds}s each, jobs={jobs}).")
    raw_dir = os.path.join(cache_dir, f"openai-asr-chunks-{asr_model}-{chunk_seconds}s")
    os.makedirs(raw_dir, exist_ok=True)

    pending = []
    raw_by_index: Dict[int, Dict[str, Any]] = {}
    for index, chunk in enumerate(chunks):
        raw_path = os.path.join(raw_dir, f"{os.path.splitext(os.path.basename(chunk))[0]}.json")
        if os.path.exists(raw_path):
            log(f"Using cached OpenAI ASR chunk {os.path.basename(chunk)} ({index + 1}/{len(chunks)})")
            raw_by_index[index] = read_json_file(raw_path)
        else:
            pending.append((index, chunk, raw_path))

    def run_one(item: Any) -> Any:
        index, chunk, raw_path = item
        chunk_name = os.path.basename(chunk)
        chunk_position = f"({index + 1}/{len(chunks)})"
        log(f"Running OpenAI ASR on {chunk_name} {chunk_position}")
        try:
            raw = transcribe_audio_chunk(
                chunk,
                openai_key,
                asr_model,
                timeout_seconds,
                log=lambda message: log(f"{message} {chunk_position}"),
            )
        except Exception as exc:
            log(f"OpenAI ASR failed {chunk_name} ({index + 1}/{len(chunks)}): {exc}")
            return {
                "ok": False,
                "index": index,
                "chunk": chunk_name,
                "chunk_path": chunk,
                "raw_path": raw_path,
                "error": str(exc),
            }
        write_json_file(raw_path, raw)
        log(f"OpenAI ASR completed {chunk_name} ({index + 1}/{len(chunks)})")
        return {
            "ok": True,
            "index": index,
            "chunk": chunk_name,
            "raw": raw,
        }

    errors = []
    remaining = pending
    for pass_number in range(1, max_passes + 1):
        if not remaining:
            break
        if pass_number > 1:
            if retry_pass_delay_seconds:
                time.sleep(retry_pass_delay_seconds)
            log(
                "Retrying failed OpenAI ASR chunks "
                f"(pass {pass_number}/{max_passes}, {len(remaining)} remaining)."
            )

        errors = []
        worker_count = min(jobs, len(remaining))
        with ThreadPoolExecutor(max_workers=worker_count) as executor:
            futures = [executor.submit(run_one, item) for item in remaining]
            for future in as_completed(futures):
                item = future.result()
                if item["ok"]:
                    raw_by_index[item["index"]] = item["raw"]
                else:
                    errors.append(item)
        remaining = [(item["index"], item["chunk_path"], item["raw_path"]) for item in errors]

    if remaining:
        error_text = "; ".join(f"{item['chunk']}: {item['error']}" for item in errors)
        raise RuntimeError(
            f"OpenAI ASR completed {len(raw_by_index)}/{len(chunks)} chunks; failed: {error_text}"
        )

    chunk_results = []
    for index, chunk in enumerate(chunks):
        raw = raw_by_index[index]
        chunk_results.append(
            {
                "chunk_index": index + 1,
                "chunk": os.path.basename(chunk),
                "offset_seconds": offsets[index],
                "segments": extract_diarized_segments(raw),
            }
        )

    result = {
        "source": "openai_asr",
        "model": asr_model,
        "chunk_seconds": chunk_seconds,
        "chunks": [
            {
                "chunk": item["chunk"],
                "offset_seconds": item["offset_seconds"],
                "segment_count": len(item["segments"]),
            }
            for item in chunk_results
        ],
        "segments": merge_diarized_chunks(chunk_results),
    }
    write_json_file(result_path, result)
    return result


def format_timecode(seconds: Optional[float]) -> str:
    if seconds is None:
        return "??:??:??"
    total = max(0, int(seconds))
    hours = total // 3600
    minutes = (total % 3600) // 60
    secs = total % 60
    return f"{hours:02d}:{minutes:02d}:{secs:02d}"


def format_segments(segments: List[Dict[str, Any]]) -> str:
    lines = []
    for idx, seg in enumerate(segments, 1):
        if seg.get("start") is not None:
            lines.append(f"[{idx} @ {format_timecode(seg.get('start'))}] {seg['text']}")
        else:
            lines.append(f"[{idx}] {seg['text']}")
    return "\n".join(lines)


def write_json_file(path: str, data: Any) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
        f.write("\n")


def read_json_file(path: str) -> Any:
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def write_text_file(path: str, text: str) -> None:
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)


def make_debug_output_dir(video_id: str, title: str) -> str:
    timestamp = time.strftime("%Y%m%d-%H%M%S")
    safe_title = sanitize_filename(title)[:60]
    folder_name = f"ytranslate-debug-{video_id}-{timestamp}"
    if safe_title:
        folder_name += f"-{safe_title}"
    output_dir = os.path.join(OUTPUT_DIR, folder_name)
    os.makedirs(output_dir, exist_ok=True)
    return output_dir


def render_markdown_transcript(
    title_translated: str,
    speakers: List[Dict[str, str]],
    turns: List[Dict[str, str]],
    metadata: Optional[Dict[str, Any]] = None,
) -> str:
    lines: List[str] = [f"# {title_translated.strip() or 'Translated Transcript'}", ""]

    if metadata:
        lines.append("## Run Info")
        lines.append("")
        for key, value in metadata.items():
            lines.append(f"- **{key}**: {value}")
        lines.append("")

    if speakers:
        lines.append("## Speakers")
        lines.append("")
        for speaker in speakers:
            label_full = speaker.get("label_full") or speaker.get("label_short") or speaker.get("id", "Speaker")
            label_short = speaker.get("label_short") or speaker.get("id", "Speaker")
            lines.append(f"- **{label_short}**: {label_full}")
        lines.append("")

    lines.append("## Transcript")
    lines.append("")
    speaker_labels = {
        speaker.get("id"): speaker.get("label_short") or speaker.get("id") or "Speaker"
        for speaker in speakers
    }
    for turn in turns:
        speaker_id = turn.get("speaker_id")
        label = speaker_labels.get(speaker_id) or speaker_id or "Speaker"
        speaker = next((item for item in speakers if item.get("id") == speaker_id), {})
        text = strip_redundant_speaker_prefix(
            turn.get("text_translated") or "",
            label,
            speaker,
        )
        if not text:
            continue
        lines.append(f"**{label}:** {text}")
        lines.append("")

    return "\n".join(lines).rstrip() + "\n"


def strip_redundant_speaker_prefix(
    text: str,
    rendered_label: str,
    speaker: Optional[Dict[str, str]] = None,
) -> str:
    cleaned = (text or "").strip()
    if not cleaned:
        return ""
    labels = [
        rendered_label,
        (speaker or {}).get("label_short", ""),
        (speaker or {}).get("label_full", ""),
        (speaker or {}).get("id", ""),
    ]
    seen = set()
    for label in sorted(labels, key=lambda value: len(value or ""), reverse=True):
        label = (label or "").strip()
        key = label.lower()
        if not label or key in seen:
            continue
        seen.add(key)
        match = re.match(rf"^\s*{re.escape(label)}\s*[:：-]\s*(.+)$", cleaned, re.IGNORECASE | re.DOTALL)
        if match:
            return match.group(1).strip()
    return cleaned


def build_target_terminology_guidance(target_language: str) -> str:
    target_norm = (target_language or "").strip().lower()
    if "russian" in target_norm or "рус" in target_norm:
        return (
            "Terminology requirements for Russian:\n"
            "- Use natural Russian equivalents where possible instead of transliterated English.\n"
            "- Never invent fake Russian-looking words such as bad calques or pseudo-technical slang.\n"
            "- Prefer fully translated phrases like 'рабочее пространство Slack' over mixed English-Russian forms.\n"
            "- Do not add bracketed explanations in this pass unless the source itself includes an explanation.\n"
        )
    return (
        "Terminology requirements:\n"
        "- Do not leave specialized English phrases untranslated in the target language.\n"
        "- If keeping acronyms, translate their expanded meaning on first mention.\n"
        "- If there is no exact equivalent, use a concise explanatory bracket.\n"
    )


def turn_key(index: int) -> str:
    return f"turn_{index:04d}"


def get_turn_text_map_schema(turn_count: int) -> Dict[str, Any]:
    properties = {
        turn_key(index): {"type": "string"}
        for index in range(1, turn_count + 1)
    }
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": properties,
        "required": list(properties.keys()),
    }


def get_turn_translation_schema(turn_count: int) -> Dict[str, Any]:
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "title_translated": {"type": "string"},
            "translations": get_turn_text_map_schema(turn_count),
        },
        "required": ["title_translated", "translations"],
    }


def get_turn_cleanup_schema(turn_count: int) -> Dict[str, Any]:
    return {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "turns": get_turn_text_map_schema(turn_count),
        },
        "required": ["turns"],
    }


def align_turn_texts_by_index(
    returned_turns: List[Dict[str, Any]],
    expected_count: int,
    pass_name: str,
) -> List[str]:
    if len(returned_turns) != expected_count:
        raise RuntimeError(f"{pass_name} pass returned the wrong number of turns")

    by_index: Dict[int, str] = {}
    for position, item in enumerate(returned_turns, 1):
        if "turn_index" not in item:
            raise RuntimeError(f"{pass_name} pass returned item {position} missing turn_index")
        try:
            turn_index = int(item.get("turn_index"))
        except (TypeError, ValueError) as exc:
            raise RuntimeError(
                f"{pass_name} pass returned invalid turn_index at item {position}"
            ) from exc
        if turn_index < 1 or turn_index > expected_count:
            raise RuntimeError(
                f"{pass_name} pass returned out-of-range turn_index {turn_index}"
            )
        if turn_index in by_index:
            raise RuntimeError(f"{pass_name} pass returned duplicate turn_index {turn_index}")
        by_index[turn_index] = (item.get("text_translated") or "").strip()

    missing = [idx for idx in range(1, expected_count + 1) if idx not in by_index]
    if missing:
        raise RuntimeError(f"{pass_name} pass missing turn_index {missing[0]}")

    return [by_index[idx] for idx in range(1, expected_count + 1)]


def align_turn_texts_by_key(
    returned_turns: Dict[str, Any],
    expected_count: int,
    pass_name: str,
) -> List[str]:
    if not isinstance(returned_turns, dict):
        raise RuntimeError(f"{pass_name} pass returned turns in an invalid shape")

    expected_keys = [turn_key(index) for index in range(1, expected_count + 1)]
    extra = sorted(set(returned_turns) - set(expected_keys))
    if extra:
        raise RuntimeError(f"{pass_name} pass returned unexpected turn key {extra[0]}")

    missing = [key for key in expected_keys if key not in returned_turns]
    if missing:
        raise RuntimeError(f"{pass_name} pass missing turn key {missing[0]}")

    return [(returned_turns[key] or "").strip() for key in expected_keys]


def format_source_turns(turns: List[Dict[str, str]]) -> str:
    lines = []
    for idx, turn in enumerate(turns, 1):
        speaker_id = turn.get("speaker_id") or "speaker"
        text = (turn.get("text_source") or "").strip()
        lines.append(f"[{turn_key(idx)}] {speaker_id}: {text}")
    return "\n".join(lines)


def build_turn_translation_system_prompt(target_language: str) -> str:
    return (
        "You are a professional transcript translator. "
        "Translate already attributed dialogue turns into the target language. "
        "Preserve the meaning, tone, and order of the conversation. "
        "Do not change speaker assignment. Do not merge turns. Do not split turns. "
        "Return one translated string for each exact turn key shown in the input brackets. "
        "Do not add information that is not present in the source. "
        "Produce idiomatic, natural-sounding language. "
        "Translate business, financial, and technical jargon into natural target-language phrasing. "
        "Do not leave obvious English terms untranslated unless they are standard in the target language. "
        "Do not invent pseudo-translations, fake calques, or unnatural target-language forms. "
        "If a natural target-language rendering exists, use it directly instead of preserving the English. "
        "Translate the video title into the target language. "
        "Return only JSON that matches the provided schema."
    )


def build_turn_translation_user_prompt(
    url: str,
    title: str,
    description: str,
    target_language: str,
    speakers: List[Dict[str, str]],
    turns: List[Dict[str, str]],
    source_language_hint: Optional[str],
) -> str:
    speaker_lines = []
    for speaker in speakers:
        label_short = speaker.get("label_short") or speaker.get("id")
        label_full = speaker.get("label_full")
        if label_full:
            speaker_lines.append(f"- {speaker.get('id')}: {label_short} ({label_full})")
        else:
            speaker_lines.append(f"- {speaker.get('id')}: {label_short}")
    source_hint = f"Source language hint: {source_language_hint}\n" if source_language_hint else ""
    terminology_guidance = build_target_terminology_guidance(target_language)
    return (
        f"Video URL: {url}\n"
        f"Title: {title}\n"
        f"Description: {description}\n"
        f"Target language: {target_language}\n"
        f"{source_hint}"
        "Speakers (keep these IDs exactly):\n"
        + "\n".join(speaker_lines)
        + "\n"
        f"{terminology_guidance}\n"
        "Turns to translate (keep order; return one translated text per input turn). "
        "Use the bracketed turn keys as JSON property names under translations. "
        "Do not put speaker labels inside translated strings:\n"
        f"{format_source_turns(turns)}"
    )


def extract_response_text(response: Any) -> Optional[str]:
    if hasattr(response, "output_text") and response.output_text:
        return response.output_text

    output = getattr(response, "output", None)
    if not output:
        return None

    texts = []
    for item in output:
        content = getattr(item, "content", None)
        if not content:
            continue
        for part in content:
            if isinstance(part, dict):
                text = part.get("text")
            else:
                text = getattr(part, "text", None)
            if text:
                texts.append(text)
    return "".join(texts) if texts else None


def is_context_length_error(err: Exception) -> bool:
    msg = str(err).lower()
    return any(
        phrase in msg
        for phrase in [
            "context length",
            "maximum context",
            "too many tokens",
            "token limit",
        ]
    )

def is_request_too_large_error(err: Exception) -> bool:
    msg = str(err).lower()
    return (
        "request too large" in msg
        or "tokens per min" in msg
        or "rate_limit_exceeded" in msg and "tokens" in msg
    )

def call_openai(
    client: OpenAI,
    model: str,
    system_prompt: str,
    user_prompt: str,
    schema_name: str = "translated_transcript",
    schema: Optional[Dict[str, Any]] = None,
    temperature: Optional[float] = None,
    reasoning_effort: Optional[str] = None,
) -> Dict[str, Any]:
    if schema is None:
        raise RuntimeError("OpenAI JSON schema must be provided")
    response_format_v1 = {
        "type": "json_schema",
        "json_schema": {
            "name": schema_name,
            "strict": True,
            "schema": schema,
        },
    }
    response_format_v2 = {
        "type": "json_schema",
        "name": schema_name,
        "strict": True,
        "schema": schema,
    }

    response_kwargs: Dict[str, Any] = {
        "model": model,
        "input": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
    }
    if reasoning_effort is None:
        response_kwargs["temperature"] = (
            OPENAI_TEMPERATURE if temperature is None else temperature
        )
    else:
        response_kwargs["reasoning"] = {"effort": reasoning_effort}

    params = inspect.signature(client.responses.create).parameters
    if "response_format" in params:
        response_kwargs["response_format"] = response_format_v1
    else:
        response_kwargs["text"] = {"format": response_format_v2}

    response = client.responses.create(**response_kwargs)

    text = extract_response_text(response)
    if not text:
        raise RuntimeError("OpenAI response contained no text")

    try:
        return json.loads(text)
    except json.JSONDecodeError as exc:
        raise RuntimeError("OpenAI response was not valid JSON") from exc


def call_openai_with_retry(
    client: OpenAI,
    model: str,
    system_prompt: str,
    user_prompt: str,
    max_retries: int = 5,
    schema_name: str = "translated_transcript",
    schema: Optional[Dict[str, Any]] = None,
    temperature: Optional[float] = None,
    reasoning_effort: Optional[str] = None,
) -> Dict[str, Any]:
    delay = 1.0
    for attempt in range(max_retries):
        try:
            return call_openai(
                client,
                model,
                system_prompt,
                user_prompt,
                schema_name=schema_name,
                schema=schema,
                temperature=temperature,
                reasoning_effort=reasoning_effort,
            )
        except Exception as exc:
            is_rate_limit = isinstance(exc, getattr(openai, "RateLimitError", ()))
            if is_rate_limit and is_request_too_large_error(exc):
                raise
            is_timeout = isinstance(exc, getattr(openai, "APITimeoutError", ()))
            is_connection = isinstance(exc, getattr(openai, "APIConnectionError", ()))
            if is_rate_limit or is_timeout or is_connection:
                if attempt == max_retries - 1:
                    raise
                time.sleep(delay + random.random())
                delay *= 2
                continue
            raise


def chunk_turns_by_speaker_and_chars(
    turns: List[Dict[str, str]],
    text_field: str,
    max_chars: int,
) -> List[List[Dict[str, str]]]:
    speaker_order: List[str] = []
    chunks_by_speaker: Dict[str, List[List[Dict[str, str]]]] = {}
    chars_by_speaker: Dict[str, List[int]] = {}
    for turn in turns:
        speaker_id = turn.get("speaker_id") or "speaker"
        if speaker_id not in chunks_by_speaker:
            speaker_order.append(speaker_id)
            chunks_by_speaker[speaker_id] = [[]]
            chars_by_speaker[speaker_id] = [0]

        chunks = chunks_by_speaker[speaker_id]
        chunk_chars = chars_by_speaker[speaker_id]
        line_len = len(turn.get(text_field, "")) + len(speaker_id) + 8
        if chunks[-1] and chunk_chars[-1] + line_len > max_chars:
            chunks.append([turn])
            chunk_chars.append(line_len)
        else:
            chunks[-1].append(turn)
            chunk_chars[-1] += line_len

    ordered_chunks: List[List[Dict[str, str]]] = []
    for speaker_id in speaker_order:
        ordered_chunks.extend(chunks_by_speaker[speaker_id])
    return ordered_chunks


def chunk_source_turns(turns: List[Dict[str, str]], max_chars: int) -> List[List[Dict[str, str]]]:
    return chunk_turns_by_speaker_and_chars(turns, "text_source", max_chars)


def translate_turn_chunk(
    client: OpenAI,
    model: str,
    url: str,
    title: str,
    description: str,
    target_language: str,
    speakers: List[Dict[str, str]],
    turns: List[Dict[str, str]],
    source_language_hint: Optional[str],
    debug_sink: Optional[List[Dict[str, Any]]] = None,
    chunk_index: int = 1,
    chunk_count: int = 1,
) -> Dict[str, Any]:
    schema = get_turn_translation_schema(len(turns))
    result = call_openai_with_retry(
        client,
        model,
        build_turn_translation_system_prompt(target_language),
        build_turn_translation_user_prompt(
            url,
            title,
            description,
            target_language,
            speakers,
            turns,
            source_language_hint,
        ),
        schema_name="translated_turn_chunk",
        schema=schema,
        temperature=OPENAI_TEMPERATURE,
    )
    if debug_sink is not None:
        debug_sink.append(
            {
                "chunk_index": chunk_index,
                "chunk_count": chunk_count,
                "turn_count": len(turns),
                "result": result,
            }
        )
    return result


def translate_attributed_turns(
    client: OpenAI,
    model: str,
    url: str,
    title: str,
    description: str,
    target_language: str,
    speakers: List[Dict[str, str]],
    turns: List[Dict[str, str]],
    source_language_hint: Optional[str],
    debug_sink: Optional[List[Dict[str, Any]]] = None,
    log: Optional[Callable[[str], None]] = None,
) -> Dict[str, Any]:
    chunks = chunk_source_turns(turns, max_chars=TURN_TEXT_PASS_MAX_CHARS)
    translated_texts_by_turn: List[str] = [""] * len(turns)
    turn_positions = {id(turn): index for index, turn in enumerate(turns)}
    title_translated = ""
    for chunk_index, chunk in enumerate(chunks, 1):
        if log:
            log(f"Translating text chunk {chunk_index}/{len(chunks)} ({len(chunk)} turns)...")
        result = translate_turn_chunk(
            client,
            model,
            url,
            title,
            description,
            target_language,
            speakers,
            chunk,
            source_language_hint,
            debug_sink=debug_sink,
            chunk_index=chunk_index,
            chunk_count=len(chunks),
        )
        if not title_translated:
            title_translated = result.get("title_translated", "")
        translated_texts = align_turn_texts_by_key(
            result.get("translations", {}),
            len(chunk),
            f"Translation chunk {chunk_index}",
        )
        for turn, translated_text in zip(chunk, translated_texts):
            translated_texts_by_turn[turn_positions[id(turn)]] = translated_text
        if log:
            log(f"Translated text chunk {chunk_index}/{len(chunks)}")

    return {
        "title_translated": title_translated,
        "speakers": speakers,
        "turns": [
            {
                "speaker_id": turn.get("speaker_id"),
                "text_translated": translated_text,
            }
            for turn, translated_text in zip(turns, translated_texts_by_turn)
        ],
    }


def is_russian_target_language(target_language: Optional[str]) -> bool:
    target_norm = (target_language or "").strip().lower()
    return "russian" in target_norm or "рус" in target_norm


def format_turns_for_cleanup(turns: List[Dict[str, str]]) -> str:
    lines = []
    for idx, turn in enumerate(turns, 1):
        speaker_id = turn.get("speaker_id") or "speaker"
        text = (turn.get("text_translated") or "").strip()
        lines.append(f"[{turn_key(idx)}] {speaker_id}: {text}")
    return "\n".join(lines)


def build_russian_cleanup_system_prompt() -> str:
    return (
        "You are a Russian-language copy editor cleaning up an already translated podcast transcript. "
        "Preserve the meaning, tone, and turn boundaries exactly, but improve wording so it reads like natural, competent Russian. "
        "Do not change speaker order, do not merge or split turns, and do not change which speaker says which turn. "
        "Return one revised string for each exact turn key shown in the input brackets. "
        "Your job is to fix awkward phrasing, mixed-language artifacts, bad grammar, bad glossary glosses, and clumsy brackets. "
        "Assume the reader is intelligent but not an expert in the domain. "
        "Keep helpful bracketed explanations for non-expert readers, but remove or rewrite bad ones. "
        "Add a short bracketed explanation on first mention only when a specialist term, acronym, internal model name, or industry phrase would otherwise be unclear to a non-expert reader. "
        "Do not add brackets for obvious brands, ordinary product names, or phrases that already read naturally in Russian. "
        "Do not invent fake Russian words or calques such as unnatural pseudo-technical slang. "
        "Do not leave ordinary English words or phrases inside brackets if they can be translated naturally into Russian. "
        "Do not add redundant glosses for obvious brand names or product names such as Ferrari, Gmail, or Google Calendar when the Russian phrase already reads naturally. "
        "Prefer a fully natural Russian phrase over mixed forms like 'Slack workspace [рабочее пространство Slack]'. "
        "If a bracketed gloss is useful, make it short, idiomatic, and genuinely informative. "
        "If a term's exact meaning is unclear, do not hallucinate certainty; keep the term and use a brief generic gloss only if needed. "
        "Examples of bad output to fix: 'биллинг [счетинг]', 'contagion' left in brackets, 'Ferrari [автомобили Ferrari]', "
        "'Google Calendar [календарь Google]', 'credits [лимит использования]', 'one-shotted it' left in English, "
        "'корпоративный долг [долг компаний]', 'Max 7 [вероятно, Max 7]', or ungrammatical phrases like 'ценный зернышко'. "
        "Return only JSON matching the provided schema."
    )


def build_russian_annotation_system_prompt() -> str:
    return (
        "You are a Russian-language editor adding concise glossary-style clarifications to an already translated podcast transcript. "
        "Preserve wording, meaning, tone, turn boundaries, and speaker assignment exactly. "
        "Do not merge turns, split turns, or rewrite sentences beyond minimal edits needed to insert a bracketed gloss. "
        "Return one revised string for each exact turn key shown in the input brackets. "
        "Assume the reader is intelligent but not an expert in the domain. "
        "Add a short Russian bracketed gloss on first mention only when a specialist term, acronym, metric, industry phrase, or product-specific concept would likely be unclear to a non-expert reader. "
        "Good candidates include terms like SLA, KYC, InMail, RAG, HVAC, InfiniBand, cryptowinter, and quant trader if they appear. "
        "Do not add glosses for obvious brands or ordinary product names such as Ferrari, Gmail, Google Calendar, Slack, Mac Mini, LinkedIn, Reddit, Claude, OpenAI, or Ethereum. "
        "Do not add glosses for plain concepts that already read naturally in Russian, such as corporate debt. "
        "Do not add glosses for already understandable Russian technical terms such as 'языковая модель', 'токены', 'субагент', or other phrases that are already self-explanatory in Russian context. "
        "If the main text already contains a clear Russian rendering, leave it alone instead of adding a bracket. "
        "Bracket text must be Russian only; never put English inside brackets. "
        "If a leftover English phrase can be translated naturally into Russian, prefer translating the phrase itself rather than explaining the English in brackets. "
        "Keep each gloss short, idiomatic, and genuinely informative. "
        "If you are not confident enough to write a useful gloss, leave the text unchanged. "
        "Never invent fake Russian words, pseudo-calques, or uncertain notes like 'вероятно'. "
        "Return only JSON matching the provided schema."
    )


def build_russian_cleanup_user_prompt(
    title_translated: str,
    turns: List[Dict[str, str]],
) -> str:
    return (
        f"Transcript title in Russian: {title_translated}\n"
        "Rewrite each turn below into better Russian while preserving meaning exactly. "
        "Keep the same number of turns and the same order. "
        "Use the bracketed turn keys as JSON property names under turns. "
        "Return only revised text values; do not include speaker labels inside the text.\n\n"
        "Turns:\n"
        f"{format_turns_for_cleanup(turns)}"
    )


def build_russian_annotation_user_prompt(
    title_translated: str,
    turns: List[Dict[str, str]],
) -> str:
    return (
        f"Transcript title in Russian: {title_translated}\n"
        "Review each turn below. Keep the same number of turns and the same order. "
        "Only add concise bracketed glosses where they are genuinely helpful for a non-expert reader. "
        "Use the bracketed turn keys as JSON property names under turns. "
        "Return only revised text values; do not include speaker labels inside the text.\n\n"
        "Turns:\n"
        f"{format_turns_for_cleanup(turns)}"
    )


def chunk_turns_by_chars(turns: List[Dict[str, str]], max_chars: int) -> List[List[Dict[str, str]]]:
    return chunk_turns_by_speaker_and_chars(turns, "text_translated", max_chars)


def cleanup_russian_turn_chunk(
    client: OpenAI,
    model: str,
    title_translated: str,
    turns: List[Dict[str, str]],
    chunk_index: int = 1,
    chunk_count: int = 1,
    debug_sink: Optional[List[Dict[str, Any]]] = None,
) -> List[str]:
    system_prompt = build_russian_cleanup_system_prompt()
    user_prompt = build_russian_cleanup_user_prompt(title_translated, turns)
    schema = get_turn_cleanup_schema(len(turns))
    result = call_openai_with_retry(
        client,
        model,
        system_prompt,
        user_prompt,
        schema_name="cleaned_transcript_turns",
        schema=schema,
        temperature=OPENAI_CLEANUP_TEMPERATURE,
    )
    if debug_sink is not None:
        debug_sink.append(
            {
                "chunk_index": chunk_index,
                "chunk_count": chunk_count,
                "turn_count": len(turns),
                "result": result,
            }
        )
    return align_turn_texts_by_key(result.get("turns", {}), len(turns), "Cleanup")


def cleanup_russian_turns(
    client: OpenAI,
    model: str,
    title_translated: str,
    turns: List[Dict[str, str]],
    debug_sink: Optional[List[Dict[str, Any]]] = None,
    log: Optional[Callable[[str], None]] = None,
) -> List[Dict[str, str]]:
    if not turns:
        return turns

    cleaned_texts = [""] * len(turns)
    turn_positions = {id(turn): index for index, turn in enumerate(turns)}
    chunks = chunk_turns_by_chars(turns, max_chars=TURN_TEXT_PASS_MAX_CHARS)
    for chunk_index, chunk in enumerate(chunks, 1):
        if log:
            log(f"Cleaning Russian text chunk {chunk_index}/{len(chunks)} ({len(chunk)} turns)...")
        chunk_texts = cleanup_russian_turn_chunk(
            client,
            model,
            title_translated,
            chunk,
            chunk_index=chunk_index,
            chunk_count=len(chunks),
            debug_sink=debug_sink,
        )
        for turn, cleaned_text in zip(chunk, chunk_texts):
            cleaned_texts[turn_positions[id(turn)]] = cleaned_text
        if log:
            log(f"Cleaned Russian text chunk {chunk_index}/{len(chunks)}")

    if len(cleaned_texts) != len(turns):
        raise RuntimeError("Cleanup pass returned the wrong number of turns")

    cleaned_turns: List[Dict[str, str]] = []
    for turn, cleaned_text in zip(turns, cleaned_texts):
        cleaned_turn = dict(turn)
        cleaned_turn["text_translated"] = cleaned_text or (turn.get("text_translated") or "")
        cleaned_turns.append(cleaned_turn)
    return cleaned_turns


def annotate_russian_turn_chunk(
    client: OpenAI,
    model: str,
    title_translated: str,
    turns: List[Dict[str, str]],
    chunk_index: int = 1,
    chunk_count: int = 1,
    debug_sink: Optional[List[Dict[str, Any]]] = None,
) -> List[str]:
    schema = get_turn_cleanup_schema(len(turns))
    result = call_openai_with_retry(
        client,
        model,
        build_russian_annotation_system_prompt(),
        build_russian_annotation_user_prompt(title_translated, turns),
        schema_name="annotated_transcript_turns",
        schema=schema,
        temperature=OPENAI_ANNOTATION_TEMPERATURE,
    )
    if debug_sink is not None:
        debug_sink.append(
            {
                "chunk_index": chunk_index,
                "chunk_count": chunk_count,
                "turn_count": len(turns),
                "result": result,
            }
        )
    return align_turn_texts_by_key(result.get("turns", {}), len(turns), "Annotation")


def annotate_russian_turns(
    client: OpenAI,
    model: str,
    title_translated: str,
    turns: List[Dict[str, str]],
    debug_sink: Optional[List[Dict[str, Any]]] = None,
    log: Optional[Callable[[str], None]] = None,
) -> List[Dict[str, str]]:
    if not turns:
        return turns

    annotated_texts = [""] * len(turns)
    turn_positions = {id(turn): index for index, turn in enumerate(turns)}
    chunks = chunk_turns_by_chars(turns, max_chars=TURN_TEXT_PASS_MAX_CHARS)
    for chunk_index, chunk in enumerate(chunks, 1):
        if log:
            log(f"Annotating glossary chunk {chunk_index}/{len(chunks)} ({len(chunk)} turns)...")
        chunk_texts = annotate_russian_turn_chunk(
            client,
            model,
            title_translated,
            chunk,
            chunk_index=chunk_index,
            chunk_count=len(chunks),
            debug_sink=debug_sink,
        )
        for turn, annotated_text in zip(chunk, chunk_texts):
            annotated_texts[turn_positions[id(turn)]] = annotated_text
        if log:
            log(f"Annotated glossary chunk {chunk_index}/{len(chunks)}")

    if len(annotated_texts) != len(turns):
        raise RuntimeError("Annotation pass returned the wrong number of turns")

    annotated_turns: List[Dict[str, str]] = []
    for turn, annotated_text in zip(turns, annotated_texts):
        annotated_turn = dict(turn)
        annotated_turn["text_translated"] = annotated_text or (turn.get("text_translated") or "")
        annotated_turns.append(annotated_turn)
    return annotated_turns


def sanitize_filename(value: str) -> str:
    normalized = unicodedata.normalize("NFKD", value)
    ascii_text = normalized.encode("ascii", "ignore").decode("ascii")
    ascii_text = re.sub(r"[^A-Za-z0-9\-_. ]+", "", ascii_text)
    ascii_text = re.sub(r"\s+", " ", ascii_text).strip()
    ascii_text = ascii_text.replace(" ", "-")
    ascii_text = re.sub(r"-+", "-", ascii_text)
    return ascii_text or "video"


def render_docx(
    title_translated: str,
    speakers: List[Dict[str, str]],
    turns: List[Dict[str, str]],
    output_path: str,
) -> None:
    doc = Document()
    doc.styles["Normal"].font.name = DOCX_FONT_NAME
    doc.styles["Normal"].font.size = DOCX_FONT_SIZE
    doc.styles["List Bullet"].font.name = DOCX_FONT_NAME
    doc.styles["List Bullet"].font.size = DOCX_FONT_SIZE
    doc.styles["Heading 1"].font.name = DOCX_FONT_NAME
    doc.styles["Heading 1"].font.size = DOCX_HEADING_FONT_SIZE
    doc.add_heading(title_translated.strip() or "Translated Transcript", level=1)

    if speakers:
        for speaker in speakers:
            label_full = speaker.get("label_full") or speaker.get("label_short") or speaker.get("id", "Speaker")
            doc.add_paragraph(label_full, style="List Bullet")

    for turn in turns:
        speaker_id = turn.get("speaker_id")
        label = None
        speaker = None
        for sp in speakers:
            if sp.get("id") == speaker_id:
                label = sp.get("label_short") or sp.get("id")
                speaker = sp
                break
        label = label or speaker_id or "Speaker"
        text = strip_redundant_speaker_prefix(
            turn.get("text_translated") or "",
            label,
            speaker,
        )
        if not text:
            continue
        para = doc.add_paragraph()
        run = para.add_run(f"{label}: ")
        run.bold = True
        para.add_run(text)

    doc.save(output_path)


def render_source_pdf(
    title: str,
    speakers: List[Dict[str, str]],
    turns: List[Dict[str, str]],
    output_dir: str,
) -> str:
    os.makedirs(output_dir, exist_ok=True)
    base_name = f"{sanitize_filename(title)}_EN"
    temporary_docx_path = os.path.join(output_dir, f"{base_name}.docx")
    source_turns = [
        {
            **turn,
            "text_translated": turn.get("text_source") or "",
        }
        for turn in turns
    ]
    try:
        render_docx(
            title,
            speakers,
            source_turns,
            temporary_docx_path,
        )
        return convert_docx_to_pdf(temporary_docx_path)
    finally:
        if os.path.exists(temporary_docx_path):
            os.remove(temporary_docx_path)


def convert_docx_to_pdf(docx_path: str) -> str:
    output_dir = os.path.dirname(docx_path) or os.getcwd()
    base_name = os.path.splitext(os.path.basename(docx_path))[0]
    expected_pdf_path = os.path.join(output_dir, f"{base_name}.pdf")

    candidate_libreoffice_paths = [
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice.bin",
    ]

    converters = [
        (shutil.which("soffice"), [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            docx_path,
        ]),
        (next((p for p in candidate_libreoffice_paths if os.access(p, os.X_OK)), None), [
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            docx_path,
        ]),
        (shutil.which("libreoffice"), [
            "libreoffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            output_dir,
            docx_path,
        ]),
    ]

    for exe, cmd in converters:
        if not exe:
            continue
        result = subprocess.run(
            cmd,
            cwd=output_dir,
            capture_output=True,
            text=True,
        )
        if result.returncode == 0 and os.path.exists(expected_pdf_path):
            return expected_pdf_path

    try:
        from docx2pdf import convert
    except Exception:
        raise RuntimeError(
            "PDF conversion is unavailable. Install one of:\n"
            "- LibreOffice (soffice/libreoffice in PATH, or a standard macOS install at /Applications/LibreOffice.app), or\n"
            "- Python package docx2pdf (pip install docx2pdf)."
        )

    convert(docx_path, expected_pdf_path)
    if not os.path.exists(expected_pdf_path):
        raise RuntimeError("docx2pdf did not generate the expected PDF file.")
    return expected_pdf_path


def send_completion_notification(message: str, title: str = "ytranslate") -> None:
    safe_title = title.replace("\\", "\\\\").replace('"', '\\"')
    safe_message = message.replace("\\", "\\\\").replace('"', '\\"')
    script = f'display notification "{safe_message}" with title "{safe_title}"'
    try:
        subprocess.run(
            ["osascript", "-e", script],
            capture_output=True,
            text=True,
            check=False,
        )
    except FileNotFoundError:
        return
    except Exception as exc:
        print(f"Failed to send macOS notification: {exc}", file=sys.stderr)

def sample_docx_payload(target_language: str) -> Dict[str, Any]:
    return {
        "title_translated": f"Пример перевода ({target_language})",
        "speakers": [
            {
                "id": "Host",
                "label_short": "Host",
                "label_full": "Alex Kantrowitz — Host, Big Technology Podcast",
            },
            {
                "id": "Guest",
                "label_short": "Guest",
                "label_full": "Demis Hassabis — CEO, Google DeepMind",
            },
        ],
        "turns": [
            {
                "speaker_id": "Host",
                "text_translated": "Это тестовый фрагмент для проверки генерации DOCX.",
            },
            {
                "speaker_id": "Guest",
                "text_translated": "Отлично. Убедимся, что перенос строк и шрифты работают.",
            },
        ],
    }


def run_sample_generation(
    target_language: Optional[str],
    log: Callable[[str], None] = print,
) -> Dict[str, Any]:
    resolved_target_language = resolve_target_language(target_language)
    output_dir = OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    filename = "sample.docx"
    output_path = os.path.join(output_dir, filename)
    sample = sample_docx_payload(resolved_target_language)
    render_docx(
        sample.get("title_translated", "Sample"),
        sample.get("speakers", []),
        sample.get("turns", []),
        output_path,
    )
    log(f"Saved sample DOCX to {output_path}")
    try:
        pdf_path = convert_docx_to_pdf(output_path)
        log(f"Saved sample PDF to {pdf_path}")
    except Exception as exc:
        raise RuntimeError(f"Failed to generate sample PDF: {exc}") from exc
    output_files = [output_path, pdf_path]
    send_completion_notification(
        "Sample conversion finished: " + ", ".join(os.path.basename(p) for p in output_files)
    )
    return {
        "target_language": resolved_target_language,
        "docx_path": output_path,
        "pdf_path": pdf_path,
        "output_files": output_files,
    }


def run_translation_job(
    url: str,
    target_language: Optional[str] = None,
    debug: bool = False,
    log: Callable[[str], None] = print,
) -> Dict[str, Any]:
    load_project_env()
    canonical_url = canonicalize_youtube_url(url)
    if not canonical_url:
        raise RuntimeError("Could not extract video ID from URL")

    resolved_target_language = resolve_target_language(target_language)
    openai_key = os.getenv("OPENAI_API_KEY")
    if not openai_key:
        raise RuntimeError("OPENAI_API_KEY is not set")

    youtube_key = os.getenv("YOUTUBE_API_KEY")
    if not youtube_key:
        raise RuntimeError("YOUTUBE_API_KEY is not set")

    video_id = extract_video_id(canonical_url)
    if not video_id:
        raise RuntimeError("Could not extract video ID from URL")

    log(f"Received URL: {canonical_url}")
    log(f"Target language: {resolved_target_language}")

    log("Fetching metadata...")
    metadata = fetch_video_metadata(video_id, youtube_key)
    title = metadata.get("title") or "Untitled"
    description = metadata.get("description", "")
    source_language_hint = metadata.get("defaultAudioLanguage") or metadata.get("defaultLanguage")
    debug_dir = make_debug_output_dir(video_id, title) if debug else None
    speaker_pass_debug: List[Dict[str, Any]] = []
    translation_pass_debug: List[Dict[str, Any]] = []
    cleanup_pass_debug: List[Dict[str, Any]] = []
    annotation_pass_debug: List[Dict[str, Any]] = []

    client = OpenAI(api_key=openai_key, timeout=OPENAI_TIMEOUT_SECONDS)
    model = os.getenv("OPENAI_MODEL", DEFAULT_MODEL)

    log("Checking YouTube transcript quality...")
    preferred_langs = [
        metadata.get("defaultAudioLanguage"),
        metadata.get("defaultLanguage"),
    ]
    transcript_info: Optional[Dict[str, Any]] = None
    youtube_segments: List[Dict[str, Any]] = []
    try:
        transcript_info = fetch_transcript(video_id, preferred_langs)
        youtube_segments = normalize_segments(transcript_info.get("segments", []))
    except (TranscriptsDisabled, NoTranscriptFound):
        log("No YouTube transcript available; using OpenAI diarized ASR.")
    except Exception as exc:
        log(f"Could not inspect YouTube transcript ({exc}); using OpenAI diarized ASR.")

    if debug_dir:
        write_json_file(
            os.path.join(debug_dir, "metadata.json"),
            metadata,
        )
        if transcript_info is not None:
            write_json_file(os.path.join(debug_dir, "youtube-transcript.json"), transcript_info)
            write_text_file(
                os.path.join(debug_dir, "youtube-normalized-transcript.md"),
                "# YouTube Normalized Transcript\n\n" + format_segments(youtube_segments) + "\n",
            )

    transcript_source = ""
    asr_result: Optional[Dict[str, Any]] = None
    speaker_mapping: Optional[Dict[str, Any]] = None
    voice_reconciliation_debug: Optional[Dict[str, Any]] = None
    speaker_identity_linker_debug: Optional[Dict[str, Any]] = None
    resolved_asr_segments: Optional[List[Dict[str, Any]]] = None
    speaker_identity_evidence: Optional[Dict[str, Any]] = None
    speaker_identity_contradictions: List[Dict[str, Any]] = []
    if transcript_info is not None and is_high_quality_youtube_transcript(transcript_info):
        log("Using manual speaker-labeled YouTube transcript.")
        transcript_source = "youtube_speaker_labeled"
        attributed = attributed_turns_from_labeled_segments(youtube_segments)
    else:
        if transcript_info is not None:
            log("YouTube transcript is not speaker-labeled; using OpenAI diarized ASR.")
        asr_result = transcribe_youtube_audio_with_openai(
            canonical_url,
            video_id,
            openai_key,
            log,
        )
        transcript_source = "openai_asr"
        log("Reconciling ASR chunk-local speakers into global speakers...")
        speaker_mapping = assign_global_speakers_for_diarized_segments(
            client,
            model,
            canonical_url,
            title,
            description,
            asr_result.get("segments", []),
            source_language_hint,
            debug_sink=speaker_pass_debug if debug else None,
            metadata=metadata,
        )
        model_speaker_mapping = speaker_mapping
        known_speaker_roster = infer_known_speaker_roster(metadata)
        if known_speaker_roster:
            log("Applying speaker identity evidence from metadata and dialogue handoffs...")
            speaker_identity_evidence = build_speaker_identity_evidence(
                asr_result.get("segments", []),
                known_speaker_roster,
            )
            speaker_mapping = apply_speaker_identity_evidence(
                speaker_mapping,
                speaker_identity_evidence,
            )
        speaker_overrides = load_speaker_mapping_overrides(video_id)
        if speaker_overrides:
            log("Applying speaker mapping overrides...")
            speaker_mapping = apply_speaker_mapping_overrides(speaker_mapping, speaker_overrides)
        cache_dir = get_video_cache_dir(video_id)
        write_json_file(os.path.join(cache_dir, "speaker-mapping-model.json"), model_speaker_mapping)
        if speaker_identity_evidence is not None:
            write_json_file(
                os.path.join(cache_dir, "speaker-identity-evidence.json"),
                serialize_speaker_identity_evidence(speaker_identity_evidence),
            )
        write_json_file(os.path.join(cache_dir, "speaker-mapping-effective.json"), speaker_mapping)
        log("Refining ASR speaker identities with voice matching...")
        resolved_asr_segments, voice_reconciliation_debug = reconcile_diarized_segments_with_voice(
            canonical_url,
            video_id,
            asr_result.get("segments", []),
            speaker_mapping,
            log,
        )
        resolved_asr_segments, effective_speakers, role_merge_debug = collapse_role_speaker_identities(
            resolved_asr_segments,
            speaker_mapping.get("speakers", []),
        )
        if role_merge_debug.get("merged_role_speakers"):
            log(
                "Merged role-like speaker identities into voice-matched speakers: "
                + ", ".join(
                    f"{role_id}->{target_id}"
                    for role_id, target_id in role_merge_debug["merged_role_speakers"].items()
                )
            )
        if isinstance(voice_reconciliation_debug, dict):
            voice_reconciliation_debug["role_speaker_identity_merge"] = role_merge_debug
        log("Linking stable speaker identities across the episode...")
        (
            resolved_asr_segments,
            effective_speakers,
            speaker_identity_linker_debug,
        ) = run_speaker_identity_linker(
            client,
            canonical_url,
            video_id,
            metadata,
            resolved_asr_segments,
            effective_speakers,
            log,
        )
        speaker_mapping_for_turns = dict(speaker_mapping)
        speaker_mapping_for_turns["speakers"] = effective_speakers
        attributed = attributed_turns_from_diarized_segments(
            resolved_asr_segments,
            speaker_mapping_for_turns,
        )
        speaker_identity_contradictions = find_speaker_identity_contradictions(
            attributed.get("speakers", []),
            attributed.get("turns", []),
        )
        write_json_file(
            os.path.join(cache_dir, "openai-asr-resolved-segments.json"),
            resolved_asr_segments,
        )
        write_json_file(
            os.path.join(cache_dir, "voice-speaker-reconciliation.json"),
            voice_reconciliation_debug,
        )
        write_json_file(
            os.path.join(cache_dir, "speaker-identity-linker.json"),
            speaker_identity_linker_debug,
        )
        write_json_file(os.path.join(cache_dir, "source-attributed-turns.json"), attributed)
        write_json_file(
            os.path.join(cache_dir, "speaker-identity-contradictions.json"),
            speaker_identity_contradictions,
        )
        if speaker_identity_contradictions:
            log(
                "Speaker identity contradiction check found "
                f"{len(speaker_identity_contradictions)} possible issue(s)."
            )

    if not attributed.get("turns"):
        raise RuntimeError("Transcript source produced no attributed turns.")

    if debug_dir:
        write_json_file(os.path.join(debug_dir, "source-attributed-turns.json"), attributed)
        if asr_result is not None:
            write_json_file(os.path.join(debug_dir, "openai-asr.json"), asr_result)
        if resolved_asr_segments is not None:
            write_json_file(
                os.path.join(debug_dir, "openai-asr-resolved-segments.json"),
                resolved_asr_segments,
            )
        if speaker_mapping is not None:
            write_json_file(os.path.join(debug_dir, "speaker-mapping.json"), speaker_mapping)
        if speaker_identity_evidence is not None:
            write_json_file(
                os.path.join(debug_dir, "speaker-identity-evidence.json"),
                serialize_speaker_identity_evidence(speaker_identity_evidence),
            )
        if voice_reconciliation_debug is not None:
            write_json_file(
                os.path.join(debug_dir, "voice-speaker-reconciliation.json"),
                voice_reconciliation_debug,
            )
        if speaker_identity_linker_debug is not None:
            write_json_file(
                os.path.join(debug_dir, "speaker-identity-linker.json"),
                speaker_identity_linker_debug,
            )
        if speaker_identity_contradictions:
            write_json_file(
                os.path.join(debug_dir, "speaker-identity-contradictions.json"),
                speaker_identity_contradictions,
            )
        if speaker_pass_debug:
            for idx, item in enumerate(speaker_pass_debug, 1):
                write_json_file(
                    os.path.join(debug_dir, f"speaker-mapping-pass-{idx:02d}.json"),
                    item,
                )
        log(f"Wrote transcript debug artifacts to {debug_dir}")

    english_pdf_path: Optional[str] = None
    if not debug_dir:
        log("Rendering English source transcript PDF...")
        try:
            english_pdf_path = render_source_pdf(
                title,
                attributed.get("speakers", []),
                attributed.get("turns", []),
                OUTPUT_DIR,
            )
            log(f"Saved English transcript PDF to {english_pdf_path}")
        except Exception as exc:
            raise RuntimeError(f"Failed to generate English transcript PDF: {exc}") from exc

    log("Translating attributed turns...")
    result = translate_attributed_turns(
        client,
        model,
        canonical_url,
        title,
        description,
        resolved_target_language,
        attributed.get("speakers", []),
        attributed.get("turns", []),
        source_language_hint,
        debug_sink=translation_pass_debug if debug else None,
        log=log,
    )

    if debug_dir:
        if translation_pass_debug:
            for idx, item in enumerate(translation_pass_debug, 1):
                write_json_file(
                    os.path.join(debug_dir, f"translation-pass-{idx:02d}.json"),
                    item,
                )
        else:
            write_json_file(os.path.join(debug_dir, "translation-pass.json"), result)

    if is_russian_target_language(resolved_target_language):
        log("Polishing Russian wording and glossary explanations...")
        result["turns"] = cleanup_russian_turns(
            client,
            model,
            result.get("title_translated", "").strip() or title,
            result.get("turns", []),
            debug_sink=cleanup_pass_debug if debug else None,
            log=log,
        )

    if debug_dir and cleanup_pass_debug:
        for idx, item in enumerate(cleanup_pass_debug, 1):
            write_json_file(
                os.path.join(debug_dir, f"cleanup-pass-{idx:02d}.json"),
                item,
            )

    if is_russian_target_language(resolved_target_language):
        log("Adding targeted glossary annotations for non-expert readers...")
        result["turns"] = annotate_russian_turns(
            client,
            model,
            result.get("title_translated", "").strip() or title,
            result.get("turns", []),
            debug_sink=annotation_pass_debug if debug else None,
            log=log,
        )

    if debug_dir and annotation_pass_debug:
        for idx, item in enumerate(annotation_pass_debug, 1):
            write_json_file(
                os.path.join(debug_dir, f"annotation-pass-{idx:02d}.json"),
                item,
            )

    title_translated = result.get("title_translated", "").strip() or title
    if debug_dir:
        write_json_file(os.path.join(debug_dir, "final.json"), result)
        final_md_path = os.path.join(debug_dir, "final.md")
        final_md = render_markdown_transcript(
            title_translated,
            result.get("speakers", []),
            result.get("turns", []),
            metadata={
                "url": canonical_url,
                "video_id": video_id,
                "target_language": resolved_target_language,
                "model": model,
                "transcript_source": transcript_source,
                "asr_model": (asr_result or {}).get("model", ""),
                "asr_chunk_seconds": (asr_result or {}).get("chunk_seconds", ""),
                "timeout_seconds": OPENAI_TIMEOUT_SECONDS,
                "temperature": OPENAI_TEMPERATURE,
                "cleanup_temperature": OPENAI_CLEANUP_TEMPERATURE,
                "annotation_temperature": OPENAI_ANNOTATION_TEMPERATURE,
                "source_language_hint": source_language_hint or "",
                "cleanup_ran": is_russian_target_language(resolved_target_language),
            },
        )
        write_text_file(final_md_path, final_md)
        log(f"Saved debug Markdown to {final_md_path}")
        log(f"Finished generating debug artifacts for {canonical_url}")
        return {
            "url": canonical_url,
            "video_id": video_id,
            "title": title,
            "title_translated": title_translated,
            "target_language": resolved_target_language,
            "debug_dir": debug_dir,
            "final_md_path": final_md_path,
            "output_files": [final_md_path],
        }

    output_dir = OUTPUT_DIR
    os.makedirs(output_dir, exist_ok=True)
    filename = f"{sanitize_filename(title)}.docx"
    output_path = os.path.join(output_dir, filename)

    render_docx(
        title_translated,
        result.get("speakers", []),
        result.get("turns", []),
        output_path,
    )
    log(f"Saved translated transcript to {output_path}")
    try:
        pdf_path = convert_docx_to_pdf(output_path)
        log(f"Saved translated transcript PDF to {pdf_path}")
    except Exception as exc:
        raise RuntimeError(f"Failed to generate transcript PDF: {exc}") from exc

    output_files = [output_path, pdf_path]
    if english_pdf_path:
        output_files.append(english_pdf_path)
    send_completion_notification(
        "Translation completed: " + ", ".join(os.path.basename(p) for p in output_files)
    )
    log(f"Finished generating files for {canonical_url}")
    return {
        "url": canonical_url,
        "video_id": video_id,
        "title": title,
        "title_translated": title_translated,
        "target_language": resolved_target_language,
        "docx_path": output_path,
        "pdf_path": pdf_path,
        "english_pdf_path": english_pdf_path,
        "output_files": output_files,
    }


def main() -> int:
    args = parse_args()
    try:
        if args.docx_test:
            if args.debug:
                raise RuntimeError("--debug is not supported together with --docx-test")
            load_project_env()
            run_sample_generation(args.target_language)
        else:
            run_translation_job(args.url, args.target_language, debug=args.debug)
    except Exception as exc:
        print(str(exc), file=sys.stderr)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
